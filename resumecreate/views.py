from django.shortcuts import render,redirect,get_object_or_404
from .models import Resume, UserProfile
from django.contrib.auth.models import User
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login, logout
from django.contrib.auth.views import LoginView, LogoutView
from django.contrib.auth.decorators import login_required
from django.urls import reverse_lazy
from django.shortcuts import redirect

from .models import UserProfile
from django import forms 


from django.http import HttpResponse
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
import json
from .services import OpenRouterService  # Import the service created
from django.template.loader import render_to_string

from django.views.decorators.http import require_POST

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
# from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, ListFlowable


# from reportlab.platypus import (
#     SimpleDocTemplate,
#     Paragraph,
#     Spacer,
#     PageBreak,
#     ListFlowable,
#     ListItem
# )

from reportlab.lib.pagesizes import letter
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.units import inch

from django.template.loader import render_to_string
from playwright.sync_api import sync_playwright
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors



# LOGIN VIEW
class CustomLoginView(LoginView):
    template_name = 'frontend/login.html'
    redirect_authenticated_user = True
    
    def get_success_url(self):
        return reverse_lazy('resumecreate:generate_cv')
    
    
# LOGOUT VIEW - ALTERNATIVE (Method 2 - Function-based - MORE RELIABLE)
def custom_logout(request):
    """
    Custom logout view that properly logs out and redirects.
    This is more reliable than the class-based view.
    """
    logout(request)
    return redirect('resumecreate:index')
    
    
# HOME VIEW
def home(request):
    """Render base.html as home/landing page"""
    context = {}
    return render(request, 'frontend/base.html', context)

# Admin dashboard view (renders base with all partials)
def admin_dashboard(request):
    registered_users = User.objects.select_related('profile').all()[:10]  # Joins profile
    submitted_resumes = Resume.objects.filter(status='submitted').order_by('-created_at')[:10]
    
    context = {
        'registered_users': registered_users, 
        'submitted_resumes': submitted_resumes,  
        'active_page': 'dashboard',
    }
    return render(request, 'admin/base_admin.html', context)

# View for registered users page (includes sidebar and user partial)
def admin_registered_users(request):
    registered_users = User.objects.select_related('profile').order_by('-date_joined')
    context = {
        'registered_users': registered_users,
        'active_page': 'users',  # Highlight in sidebar
    }
    return render(request, 'admin/base_admin.html', context)

# View for submitted resumes page
def admin_submitted_resumes(request):
    submitted_resumes = Resume.objects.filter(status='submitted').order_by('-created_at')
    context = {
        'submitted_resumes': submitted_resumes,
        'active_page': 'resumes',
    }
    return render(request, 'admin/base_admin.html', context)

# Detail view for a specific resume 
def admin_resume_detail(request, resume_id):
    # Find dummy resume or 404
    resume = next((r for r in DUMMY_RESUMES if r['id'] == resume_id), None)
    if not resume:
        # In real app, use get_object_or_404; here, dummy 404 context
        context = {'error': 'Resume not found'}
        return render(request, 'admin/partials/_resume_detail.html', context)
    
    context = {
        'resume': resume,
        'active_page': 'resumes',
    }
    return render(request, 'admin/partials/_resume_detail.html', context)  # Adjust template if needed

#E##########
class CustomUserCreationForm(UserCreationForm):
    email = forms.EmailField(required=True, label="Email Address")  # Add email field

    class Meta:
        model = User  # Django's User model
        fields = ('username', 'email', 'password1', 'password2')  # Include email

    def save(self, commit=True):
        user = super().save(commit=False)
        user.email = self.cleaned_data['email']
        if commit:
            user.save()
        return user

def register(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            UserProfile.objects.get_or_create(user=user)  # Auto-profile
            login(request, user)
            return redirect('resumecreate:generate_cv')  # Or your builder URL
    else:
        form = CustomUserCreationForm()
    return render(request, 'frontend/signup.html', {'form': form})


# def generate_cv(request): 
#     return render(request, 'frontend/generate_cv.html', {}) 


#reset conversation for resume builder
@login_required
@require_POST
def reset_resume_conversation(request):
    """
    Resets the current AI resume conversation and draft resume sections.
    """
    service = OpenRouterService()
    success = service.reset_user_conversation(request.user, request=request)

    if success:
        return JsonResponse({'status': 'success', 'message': 'Conversation reset'})
    else:
        return JsonResponse({'status': 'error', 'message': 'Failed to reset conversation'})





#generate cv view
@login_required
def generate_cv(request):
    """
    Main CV generation view with AI chat integration.
    COMPLETELY SEPARATE from cover letter builder.
    """
    user_initials = request.user.username[:2].upper()
    
    # IMPORTANT: Exclude cover letters from resume query
    resume = Resume.objects.filter(
        user=request.user, 
        status='draft'
    ).exclude(
        title__exact='Cover Letter Draft'  # Exclude cover letters
    ).order_by('-created_at').first()
    
    # Create new resume if none exists
    if not resume:
        resume = Resume.objects.create(
            user=request.user, 
            status='draft', 
            title='New Resume',
            sections={
                'type': 'resume',  # Mark as resume
                'conversation': [],
                'progress': {'current_step': 1}
            }
        )
    
    # Ensure sections has required structure
    if 'type' not in resume.sections:
        resume.sections['type'] = 'resume'
    if 'conversation' not in resume.sections:
        resume.sections['conversation'] = []
    if 'progress' not in resume.sections:
        resume.sections['progress'] = {'current_step': 1}
        resume.save()
    
    # Handle reset/clear conversation
    if request.GET.get('reset') == 'true' or request.POST.get('reset') == 'true':
        resume.sections = {
            'type': 'resume',
            'conversation': [],
            'progress': {'current_step': 1}
        }
        resume.save()
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': True, 'message': 'Conversation reset!'})
        return redirect('resumecreate:generate_cv')

    conversation_history = resume.sections.get('conversation', [])

    # Initialize progress data
    progress_data = {
        'current_step': resume.sections.get('progress', {}).get('current_step', 1),
        'total_steps': 6,
        'steps': [
            {"id": 1, "title": "Personal Info", "description": "Tell me your name, contact info, and location"},
            {"id": 2, "title": "Professional Summary", "description": "What's your career objective or professional summary?"},
            {"id": 3, "title": "Work Experience", "description": "Let's detail your work history and achievements"},
            {"id": 4, "title": "Education", "description": "Share your educational background"},
            {"id": 5, "title": "Skills & Certifications", "description": "What are your key skills and certifications?"},
            {"id": 6, "title": "Final Review", "description": "Review and finalize your resume"}
        ],
        'percentage': 16
    }

    if conversation_history:
        user_messages = [msg for msg in conversation_history if msg.get('role') == 'user']
        step_count = len(user_messages)
        current_step = min(step_count, 6) if step_count > 0 else 1
        progress_data['current_step'] = current_step
        progress_data['percentage'] = int((current_step / 6) * 100)
        resume.sections.setdefault('progress', {})['current_step'] = current_step
        resume.save()

    templates = [
        {'id': 1, 'image': 'frontend/images/1.png', 'name': 'Modern'},
        {'id': 2, 'image': 'frontend/images/2.png', 'name': 'Classic'},
        {'id': 3, 'image': 'frontend/images/3.png', 'name': 'Creative'},
        {'id': 4, 'image': 'frontend/images/4.png', 'name': 'Minimal'},
        {'id': 5, 'image': 'frontend/images/5.png', 'name': 'Professional'},
        {'id': 6, 'image': 'frontend/images/6.png', 'name': 'Tech'},
    ]

    ai_messages = []
    if not conversation_history or len(conversation_history) == 0:
        ai_messages = [{
            'avatar': 'AI',
            'text': 'Nice choice! Would you like to improve your current resume, or start from scratch?',
            'is_user': False,
            'show_actions': True
        }]
    else:
        for msg in conversation_history:
            if msg['role'] != 'system':
                ai_messages.append({
                    'avatar': user_initials if msg['role'] == 'user' else 'AI',
                    'text': msg['content'],
                    'is_user': msg['role'] == 'user',
                    'show_actions': False
                })

    if request.method == 'POST':
        try:
            ai_service = OpenRouterService()
            user_context = {
                'username': request.user.username,
                'email': request.user.email,
                'profile': {
                    'phone': getattr(request.user.profile, 'phone', ''),
                    'country': getattr(request.user.profile, 'country', ''),
                    'bio': getattr(request.user.profile, 'bio', '')
                } if hasattr(request.user, 'profile') else {},
                'current_step': progress_data['current_step'],
                'current_step_title': progress_data['steps'][progress_data['current_step']-1]['title']
            }

            # Handle template selection
            if 'template_id' in request.POST:
                template_id = request.POST.get('template_id')
                resume.sections['template_id'] = template_id
                resume.sections.setdefault('conversation', []).append({
                    'role': 'system',
                    'content': f'User selected template {template_id}'
                })
                resume.sections['type'] = 'resume'  # Ensure type is set
                resume.save()
                return JsonResponse({'success': True, 'message': 'Template selected successfully'})

            # Handle action buttons
            if 'action' in request.POST:
                action = request.POST.get('action')
                if action == 'improve':
                    user_message = "I'd like to improve my existing resume"
                elif action == 'fresh':
                    user_message = "I want to start from scratch"
                    progress_data['current_step'] = 1
                    progress_data['percentage'] = 16
                    resume.sections['progress']['current_step'] = 1
                else:
                    user_message = action

                conversation_history.append({'role': 'user', 'content': user_message})

                current_section = progress_data['steps'][progress_data['current_step']-1]['title'].lower()
                ai_response = ai_service.generate_resume_section(
                    request.user,
                    section_type=current_section,
                    user_data={"user_input": user_message}
                )

                conversation_history.append({'role': 'assistant', 'content': ai_response})

                # Advance progress
                if action != 'fresh':
                    progress_data['current_step'] = min(progress_data['current_step'] + 1, 6)
                progress_data['percentage'] = int((progress_data['current_step'] / 6) * 100)

                # Save updates
                resume.sections['conversation'] = conversation_history
                resume.sections['progress']['current_step'] = progress_data['current_step']
                resume.sections['type'] = 'resume'  # Ensure type is set
                resume.save()

                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': True,
                        'response': ai_response,
                        'avatar': 'AI',
                        'progress_data': progress_data
                    })

            # Handle chat input
            if 'chat_input' in request.POST:
                user_input = request.POST.get('chat_input', '').strip()
                if user_input:
                    conversation_history.append({'role': 'user', 'content': user_input})
                    current_section = progress_data['steps'][progress_data['current_step']-1]['title'].lower()
                    ai_response = ai_service.generate_resume_section(
                        request.user,
                        section_type=current_section,
                        user_data={"user_input": user_input}
                    )

                    conversation_history.append({'role': 'assistant', 'content': ai_response})

                    # Advance step if not last
                    if progress_data['current_step'] < 6:
                        progress_data['current_step'] += 1
                        progress_data['percentage'] = int((progress_data['current_step'] / 6) * 100)

                    # Save conversation
                    resume.sections['conversation'] = conversation_history
                    resume.sections['progress']['current_step'] = progress_data['current_step']
                    resume.sections['type'] = 'resume'  # Ensure type is set
                    resume.save()

                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return JsonResponse({
                            'success': True,
                            'response': ai_response,
                            'avatar': 'AI',
                            'progress_data': progress_data
                        })
                    else:
                        return redirect('resumecreate:generate_cv')

            # Handle export request
            if 'export' in request.POST:
                export_type = request.POST.get('export')
                resume.status = 'generated'
                resume.save()
                return JsonResponse({'success': True, 'message': f'Export to {export_type} requested'})
                
        except Exception as e:
            print(f"Resume generation error: {e}")  # Debug logging
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'error': str(e)})

    context = {
        'user_initials': user_initials,
        'templates': templates,
        'ai_messages': ai_messages,
        'resume_id': resume.id,
        'has_conversation': len(conversation_history) > 0,
        'current_step': progress_data['current_step'],
        'total_steps': progress_data['total_steps'],
        'steps': progress_data['steps'],
        'progress_percentage': progress_data['percentage']
    }

    return render(request, 'frontend/generate_cv.html', context)



#reset conversation
@login_required
def reset_conversation(request):
    """
    Resets the user's conversation and draft resume.
    """
    if request.method == 'POST':
        service = OpenRouterService()
        success = service.reset_user_conversation(request.user, request=request)
        return JsonResponse({'success': success})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=400)


#generate cover
@login_required
def generate_cover(request):
    """
    Cover letter generation view with AI integration and progress tracking.
    """
    user_initials = request.user.username[:2].upper()
    
    # Get or create draft cover letter (reusing Resume model)
    # Use a specific identifier to separate cover letters from resumes
    cover_letter = Resume.objects.filter(
        user=request.user,
        status='draft',
        title__exact='Cover Letter Draft'  # Exact match to avoid confusion
    ).order_by('-created_at').first()
    
    if not cover_letter:
        cover_letter = Resume.objects.create(
            user=request.user,
            status='draft',
            title='Cover Letter Draft',
            sections={'type': 'cover_letter', 'conversation': [], 'progress': {
                'job_details': False,
                'key_skills': False,
                'experience': False,
                'motivation': False
            }}
        )
    
    # Check if user wants to start fresh
    if request.GET.get('clear') == 'true':
        if cover_letter:
            cover_letter.sections = {
                'type': 'cover_letter',
                'conversation': [],
                'progress': {
                    'job_details': False,
                    'key_skills': False,
                    'experience': False,
                    'motivation': False
                }
            }
            cover_letter.save()
        return redirect('resumecreate:generate_cover')
    
    conversation_history = cover_letter.sections.get('conversation', [])
    
    # Calculate progress
    progress = cover_letter.sections.get('progress', {
        'job_details': False,
        'key_skills': False,
        'experience': False,
        'motivation': False
    })
    
    # Calculate completion percentage
    completed_items = sum(1 for v in progress.values() if v)
    completion_percentage = int((completed_items / 4) * 100)
    
    # Initial messages
    ai_messages = []
    if not conversation_history:
        ai_messages = [{
            'avatar': 'AI',
            'text': "Hello! I'm your AI cover letter assistant. To create a compelling cover letter, I'll need some information. Let's start with: What's the job title and company name you're applying to?",
            'is_user': False
        }]
    else:
        for msg in conversation_history:
            if msg['role'] != 'system':
                ai_messages.append({
                    'avatar': user_initials if msg['role'] == 'user' else 'AI',
                    'text': msg['content'],
                    'is_user': msg['role'] == 'user'
                })
    
    if request.method == 'POST':
        # Handle file upload
        if 'resume_file' in request.FILES:
            resume_file = request.FILES['resume_file']
            # TODO: Process resume file (extract text, save to model)
            # For now, just acknowledge receipt
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True,
                    'message': f'Resume "{resume_file.name}" uploaded successfully! Now, tell me about the job you\'re applying for.'
                })
        
        # Handle chat input
        user_input = request.POST.get('chat_input', '').strip()
        
        if user_input:
            conversation_history.append({'role': 'user', 'content': user_input})
            
            try:
                ai_service = OpenRouterService()
                
                # Cover letter specific system prompt
                system_prompt = """You are an expert cover letter writer. Help users create compelling, personalized cover letters.

Your approach:
1. Ask about job details (company, position, job description)
2. Inquire about relevant skills and experience
3. Understand their motivation for the role
4. Guide them to highlight achievements that match the job requirements

Keep responses conversational, encouraging, and professional. Ask one question at a time."""

                messages = [{'role': 'system', 'content': system_prompt}]
                
                # Add context if we have progress data
                if any(progress.values()):
                    context = f"Progress so far: {progress}"
                    messages.append({'role': 'system', 'content': context})
                
                messages.extend(conversation_history[:-1])
                messages.append({'role': 'user', 'content': user_input})
                
                ai_response = ai_service._make_request(messages)
                
                conversation_history.append({'role': 'assistant', 'content': ai_response})
                
                # Update progress based on conversation
                user_input_lower = user_input.lower()
                if any(word in user_input_lower for word in ['company', 'position', 'job', 'role', 'apply']):
                    progress['job_details'] = True
                if any(word in user_input_lower for word in ['skill', 'experience', 'python', 'javascript', 'manage', 'lead']):
                    progress['key_skills'] = True
                if any(word in user_input_lower for word in ['worked', 'years', 'developed', 'managed', 'led']):
                    progress['experience'] = True
                if any(word in user_input_lower for word in ['passionate', 'excited', 'because', 'want', 'interested']):
                    progress['motivation'] = True
                
                # Save to cover letter
                cover_letter.sections['conversation'] = conversation_history
                cover_letter.sections['progress'] = progress
                cover_letter.save()
                
                # Calculate new completion
                completed_items = sum(1 for v in progress.values() if v)
                completion_percentage = int((completed_items / 4) * 100)
                
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': True,
                        'response': ai_response,
                        'avatar': 'AI',
                        'progress': {
                            'job_details': progress['job_details'],
                            'key_skills': progress['key_skills'],
                            'experience': progress['experience'],
                            'motivation': progress['motivation'],
                            'percentage': completion_percentage
                        }
                    })
                else:
                    return redirect('resumecreate:generate_cover')
                    
            except Exception as e:
                error_message = f"Sorry, I encountered an error: {str(e)}"
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({'success': False, 'error': error_message})
        
        # Handle export
        if 'export' in request.POST:
            export_type = request.POST.get('export')
            cover_letter.status = 'generated'
            cover_letter.save()
            # TODO: Generate actual PDF/DOC file
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True, 
                    'message': f'Cover letter exported as {export_type}'
                })
    
    # Rebuild messages for display
    if conversation_history:
        ai_messages = []
        for msg in conversation_history:
            if msg['role'] != 'system':
                ai_messages.append({
                    'avatar': user_initials if msg['role'] == 'user' else 'AI',
                    'text': msg['content'],
                    'is_user': msg['role'] == 'user'
                })
    
    context = {
        'user_initials': user_initials,
        'ai_messages': ai_messages,
        'progress': progress,
        'completion_percentage': completion_percentage,
    }
    
    return render(request, 'frontend/generate_cover_letter.html', context)

    
#keep history for Gisele Knowles

@login_required
def user_history(request):
    user_initials = request.user.username[:2].upper() 

    resumes = Resume.objects.filter(user=request.user).order_by('-created_at')[:10]  

    context = {
        'user_initials': user_initials,
        'resumes': resumes,  # For card loop
    }

    if request.method == 'POST':
        search_query = request.POST.get('search')
        if search_query:
            resumes = resumes.filter(title__icontains=search_query)  # Filter by title
            context['resumes'] = resumes

    return render(request, 'frontend/history.html', context)


#resume detail

@login_required
def resume_detail(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)  # Secure: only own resumes
    context = {
        'resume': resume,  
    }
    return render(request, 'frontend/resume_detail.html', context)

#recommended Jobs
@login_required
def recommend_jobs_view(request):
    ai_service = OpenRouterService()
    user = request.user
    recommended_jobs = []
    current_job_title = ""

    if request.method == "POST":
        current_job_title = request.POST.get("current_job_title", "").strip()
        if current_job_title:
            # Get AI response string
            ai_response = ai_service.recommend_jobs(user, current_job_title)
            # Convert to a clean list
            recommended_jobs = [line.strip() for line in ai_response.split("\n") if line.strip()]

    return render(request, "frontend/recommended_jobs.html", {
        "recommended_jobs": recommended_jobs,
        "current_job_title": current_job_title,
        "user_initials": user.username[:2].upper(),
    })




#exporting resume to pdf, Beyonce
@login_required
def export_resume(request):
    """
    Export user's AI-generated resume as PDF, DOCX, or TXT.
    Extracts content from conversation history.
    """
    format_type = request.GET.get("format", "pdf")
    
    # Get the resume
    resume = Resume.objects.filter(
        user=request.user, 
        status='draft'
    ).exclude(
        title__exact='Cover Letter Draft'
    ).order_by('-created_at').first()

    if not resume:
        # Create a basic resume to export
        return HttpResponse(
            """
            <html>
            <body style="font-family: Arial; padding: 40px; text-align: center;">
                <h2>No Resume Found</h2>
                <p>Please create a resume first by answering the AI questions.</p>
                <p><a href="/generate-cv/" style="color: #2563EB;">← Go back to Resume Builder</a></p>
            </body>
            </html>
            """,
            status=200  # Don't show error, show friendly message
        )

    # Extract resume data from conversation
    try:
        resume_data = extract_resume_from_conversation(resume)
    except Exception as e:
        print(f"Extraction error: {e}")
        # Create minimal fallback data
        resume_data = {
            'name': request.user.get_full_name() or request.user.username,
            'email': request.user.email,
            'phone': '',
            'location': '',
            'linkedin': '',
            'summary': 'Professional with diverse experience',
            'experience': [],
            'education': [],
            'skills': ['Communication', 'Teamwork', 'Problem Solving'],
            'certifications': []
        }
    
    # Validate resume_data exists
    if not resume_data:
        return HttpResponse(
            """
            <html>
            <body style="font-family: Arial; padding: 40px; text-align: center;">
                <h2>Not Enough Information</h2>
                <p>Please complete more steps in the resume builder before exporting.</p>
                <p><a href="/generate-cv/" style="color: #2563EB;">← Continue Building Resume</a></p>
            </body>
            </html>
            """,
            status=200
        )

    # Generate the appropriate format
    try:
        if format_type == "pdf":
            return generate_pdf_resume(request.user, resume_data)
        elif format_type == "doc":
            return generate_docx_resume(request.user, resume_data)
        elif format_type == "txt":
            return generate_txt_resume(request.user, resume_data)
        else:
            return HttpResponse("Invalid format. Please choose PDF, DOC, or TXT.", status=400)
    except Exception as e:
        print(f"Export error: {e}")
        return HttpResponse(
            f"""
            <html>
            <body style="font-family: Arial; padding: 40px;">
                <h2>Export Error</h2>
                <p>There was an error generating your resume: {str(e)}</p>
                <p>Please try again or contact support if the issue persists.</p>
                <p><a href="/generate-cv/" style="color: #2563EB;">← Go back</a></p>
            </body>
            </html>
            """,
            status=200
        )
    

import re
from io import BytesIO
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors


def parse_experience_entry(exp_text):
    """Extract work experience details from text."""
    experience_entry = {
        'title': '',
        'company': '',
        'duration': '',
        'description': exp_text.strip()
    }
    
    # Extract job title
    title_pattern = r'(software developer|software engineer|full stack|backend|frontend|data scientist|data analyst|product manager|project manager|business analyst|web developer|mobile developer|devops engineer|system administrator|network engineer|security analyst|ui/ux designer|graphic designer|content writer|marketing manager|sales manager|account manager|hr manager|financial analyst|consultant|teacher|professor|researcher)'
    title_match = re.search(title_pattern, exp_text, re.IGNORECASE)
    
    if title_match:
        experience_entry['title'] = title_match.group().title()
    else:
        first_sentence = exp_text.split('.')[0]
        experience_entry['title'] = first_sentence if len(first_sentence) < 50 else "Professional"
    
    # Extract company name
    company_indicators = ['at', 'for', 'with', 'company', 'inc', 'ltd', 'corp']
    words = exp_text.split()
    for i, word in enumerate(words):
        if word.lower() in company_indicators and i + 1 < len(words):
            potential_company = ' '.join(words[i+1:min(i+4, len(words))])
            if potential_company and potential_company[0].isupper():
                experience_entry['company'] = potential_company.split('.')[0]
                break
    
    if not experience_entry['company']:
        experience_entry['company'] = "Company"
    
    # Extract duration
    duration_patterns = [
        r'(\d+)\s*(year|yr|month|mo)s?',
        r'(20\d{2})\s*-?\s*(20\d{2}|present|current)',
        r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s*20\d{2}'
    ]
    for pattern in duration_patterns:
        duration_match = re.search(pattern, exp_text, re.IGNORECASE)
        if duration_match:
            experience_entry['duration'] = duration_match.group()
            break
    
    if not experience_entry['duration']:
        experience_entry['duration'] = "Recent"
    
    return experience_entry


def parse_education_entry(response):
    """Extract education details from text."""
    education_entry = {
        'degree': '',
        'institution': '',
        'year': '',
        'details': response.strip()
    }
    
    # Extract degree
    degree_patterns = [
        r'(bachelor|master|phd|doctorate|bsc|msc|ba|ma|mba|associate)',
        r'(b\.?\s*(?:sc|a|eng|tech|com))',
        r'(m\.?\s*(?:sc|a|eng|tech|com|ba))'
    ]
    for pattern in degree_patterns:
        degree_match = re.search(pattern, response, re.IGNORECASE)
        if degree_match:
            education_entry['degree'] = degree_match.group().title()
            break
    
    if not education_entry['degree']:
        education_entry['degree'] = "Degree"
    
    # Extract institution
    institution_keywords = ['university', 'college', 'institute', 'school']
    words = response.split()
    for i, word in enumerate(words):
        if any(inst in word.lower() for inst in institution_keywords):
            start = max(0, i-2)
            end = min(len(words), i+3)
            potential_inst = ' '.join(words[start:end])
            education_entry['institution'] = potential_inst.split('.')[0]
            break
    
    if not education_entry['institution']:
        education_entry['institution'] = "University"
    
    # Extract year
    year_match = re.search(r'(19|20)\d{2}', response)
    if year_match:
        education_entry['year'] = year_match.group()
    else:
        education_entry['year'] = "Recent"
    
    return education_entry


def extract_skills(user_responses):
    """Extract skills from user responses."""
    skill_database = [
        'python', 'javascript', 'java', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift',
        'html', 'css', 'react', 'angular', 'vue', 'node', 'express', 'django', 'flask',
        'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'aws', 'azure', 'gcp', 'docker',
        'kubernetes', 'git', 'jenkins', 'ci/cd', 'agile', 'scrum', 'jira', 'figma',
        'leadership', 'communication', 'teamwork', 'problem solving', 'analytical',
        'project management', 'time management', 'critical thinking'
    ]
    
    all_text = ' '.join(user_responses).lower()
    found_skills = []
    
    for skill in skill_database:
        if skill in all_text:
            found_skills.append(skill.title())
    
    # Check last response for additional skills
    if len(user_responses) > 4:
        last_response = user_responses[-1]
        potential_skills = re.split(r'[,\n]|and', last_response)
        for skill in potential_skills:
            clean_skill = skill.strip().title()
            if 3 < len(clean_skill) < 30 and clean_skill not in found_skills:
                found_skills.append(clean_skill)
    
    # Remove duplicates and limit
    unique_skills = list(dict.fromkeys(found_skills))[:20]
    
    return unique_skills if unique_skills else ['Communication', 'Teamwork', 'Problem Solving']


def extract_resume_from_conversation(resume):
    """
    COMPREHENSIVE extraction - captures ALL conversation content properly.
    Maps conversation to resume sections based on step progression.
    ALWAYS returns valid data (never None).
    """
    conversation = resume.sections.get('conversation', [])
    
    # Initialize comprehensive resume data with defaults
    resume_data = {
        'name': resume.user.get_full_name() or resume.user.username,
        'email': resume.user.email,
        'phone': '',
        'location': '',
        'linkedin': '',
        'summary': '',
        'experience': [],
        'education': [],
        'skills': [],
        'certifications': []
    }
    
    # If no conversation at all, return defaults
    if not conversation:
        resume_data['summary'] = 'Professional seeking new opportunities'
        resume_data['skills'] = ['Communication', 'Teamwork', 'Problem Solving']
        return resume_data
    
    # Separate user responses and AI responses
    user_responses = [msg['content'] for msg in conversation if msg.get('role') == 'user']
    ai_responses = [msg['content'] for msg in conversation if msg.get('role') == 'assistant']
    
    # Remove opener sentences that do not carry real data
    openers = {'i want', 'i would like', 'can you', 'could you', 'please'}
    cleaned = []
    for r in user_responses:
        low = r.lower()
        if any(low.startswith(o) for o in openers):
            continue
        cleaned.append(r)
    user_responses = cleaned
    
    # Don't return None - always return something
    if len(user_responses) == 0:
        resume_data['summary'] = 'Professional seeking new opportunities'
        resume_data['skills'] = ['Communication', 'Teamwork', 'Problem Solving']
        return resume_data
    
    # ═══════════════════════════════════════════════════════════
    # STEP 1: CONTACT INFO (First user response)
    # ═══════════════════════════════════════════════════════════
    contact_info = user_responses[0]
    
    # Extract email
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', contact_info)
    if email_match:
        resume_data['email'] = email_match.group()
    
    # Extract phone (multiple formats)
    phone_patterns = [
        r'\b\d{10,11}\b',
        r'\b\(\d{3}\)\s*\d{3}-\d{4}\b',
        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
    ]
    for pattern in phone_patterns:
        phone_match = re.search(pattern, contact_info)
        if phone_match:
            resume_data['phone'] = phone_match.group()
            break
    
    # Extract name (first 2-4 words if they look like a name)
    words = contact_info.split()
    for i in range(2, min(5, len(words) + 1)):
        potential_name = ' '.join(words[:i])
        if '@' not in potential_name and not any(char.isdigit() for char in potential_name):
            if any(word[0].isupper() for word in words[:i]):
                resume_data['name'] = potential_name
                break
    
    # Extract location
    location_keywords = ['nigeria', 'lagos', 'abuja', 'uk', 'usa', 'canada', 'india']
    for keyword in location_keywords:
        if keyword in contact_info.lower():
            sentences = contact_info.split('.')
            for sent in sentences:
                if keyword in sent.lower():
                    resume_data['location'] = sent.strip()
                    break
            break
    
    # ═══════════════════════════════════════════════════════════
    # STEP 2: PROFESSIONAL SUMMARY (Second user response)
    # ═══════════════════════════════════════════════════════════
    if len(user_responses) > 1:
        summary_response = user_responses[1].strip()
        resume_data['summary'] = summary_response
        
        # Enhance if too short
        if len(summary_response) < 50 and len(ai_responses) > 1:
            for ai_resp in ai_responses[1:3]:
                if any(kw in ai_resp.lower() for kw in ['summary', 'objective']):
                    if 50 < len(ai_resp) < 500:
                        resume_data['summary'] = ai_resp.strip()
                        break
    
    # ═══════════════════════════════════════════════════════════
    # STEP 3: WORK EXPERIENCE (Third+ user responses)
    # ═══════════════════════════════════════════════════════════
    experience_keywords = [
        'worked', 'developer', 'engineer', 'manager', 'analyst',
        'consultant', 'designer', 'administrator', 'specialist',
        'years', 'months', 'company', 'role', 'position', 'job'
    ]
    
    if len(user_responses) > 2:
        for response in user_responses[2:]:
            if any(kw in response.lower() for kw in experience_keywords):
                experience_entry = parse_experience_entry(response)
                if experience_entry:
                    resume_data['experience'].append(experience_entry)
    
    # ═══════════════════════════════════════════════════════════
    # STEP 4: EDUCATION
    # ═══════════════════════════════════════════════════════════
    education_keywords = [
        'degree', 'bachelor', 'master', 'phd', 'bsc', 'msc',
        'university', 'college', 'graduated', 'education', 'diploma'
    ]
    
    for response in user_responses:
        if any(kw in response.lower() for kw in education_keywords):
            education_entry = parse_education_entry(response)
            if education_entry:
                resume_data['education'].append(education_entry)
    
    # ═══════════════════════════════════════════════════════════
    # STEP 5: SKILLS
    # ═══════════════════════════════════════════════════════════
    resume_data['skills'] = extract_skills(user_responses)
    
    return resume_data


def generate_pdf_resume(user, resume_data):
    """
    Generate a professional PDF resume with 'RESUME' as header.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.6*inch,
        leftMargin=0.6*inch,
        topMargin=0.5*inch,
        bottomMargin=0.4*inch,
    )

    styles = getSampleStyleSheet()

    # Define custom styles
    styles.add(ParagraphStyle(
        name='ResumeHeader',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a2540'),
        spaceAfter=20,
        spaceBefore=0,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    ))

    styles.add(ParagraphStyle(
        name='CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#1a2540'),
        spaceAfter=8,
        spaceBefore=10,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    ))

    styles.add(ParagraphStyle(
        name='CustomContact',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=20,
        textColor=colors.HexColor('#4a5568')
    ))

    styles.add(ParagraphStyle(
        name='CustomSectionHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#2563EB'),
        spaceAfter=10,
        spaceBefore=16,
        fontName='Helvetica-Bold'
    ))

    styles.add(ParagraphStyle(
        name='CustomJobTitle',
        parent=styles['Normal'],
        fontSize=12,
        fontName='Helvetica-Bold',
        spaceAfter=4,
        textColor=colors.HexColor('#1a2540')
    ))

    styles.add(ParagraphStyle(
        name='CustomCompanyDuration',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
        textColor=colors.HexColor('#4a5568'),
        fontName='Helvetica-Oblique'
    ))

    styles.add(ParagraphStyle(
        name='CustomBodyText',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=8,
        alignment=TA_JUSTIFY,
        textColor=colors.HexColor('#2d3748')
    ))

    styles.add(ParagraphStyle(
        name='CustomBulletText',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=4,
        leftIndent=20,
        textColor=colors.HexColor('#2d3748')
    ))

    story = []

    # ----------------  HEADER WITH "RESUME"  -----------------
    story.append(Paragraph("RESUME", styles['ResumeHeader']))
    story.append(Spacer(1, 0.1 * inch))
    


    # ----------------  CONTACT INFO  -----------------
    contact_parts = []
    if resume_data.get('email'):
        contact_parts.append(f"✉ {resume_data['email']}")
    if resume_data.get('phone'):
        contact_parts.append(f"☎ {resume_data['phone']}")
    if resume_data.get('location'):
        contact_parts.append(f"📍 {resume_data['location']}")
    
    if contact_parts:
        story.append(
            Paragraph(" &nbsp;&nbsp;|&nbsp;&nbsp; ".join(contact_parts),
                      styles['CustomContact'])
        )

    story.append(HRFlowable(
        width='85%',
        thickness=1,
        color=colors.HexColor('#d1d5db'),
        spaceBefore=2,
        spaceAfter=12
    ))

    # ----------------  SUMMARY  -----------------
    if resume_data.get('summary'):
        story.append(Paragraph("PROFESSIONAL SUMMARY", styles['CustomSectionHeading']))
        summary = re.sub(
            r'(what is your|tell me|what do you|let me know)',
            '',
            resume_data['summary'],
            flags=re.I
        ).strip()
        story.append(Paragraph(summary, styles['CustomBodyText']))
        story.append(Spacer(1, 0.15 * inch))

    # ----------------  EXPERIENCE  -----------------
    if resume_data.get('experience'):
        story.append(Paragraph("WORK EXPERIENCE", styles['CustomSectionHeading']))
        for exp in resume_data['experience']:
            story.append(Paragraph(exp['title'], styles['CustomJobTitle']))
            
            company_dur = exp['company']
            if exp.get('duration'):
                company_dur += f" • {exp['duration']}"
            story.append(Paragraph(company_dur, styles['CustomCompanyDuration']))

            desc = exp['description'].strip()
            if '•' in desc or '\n-' in desc:
                story.append(Paragraph(
                    desc.replace('\n', '<br/>'),
                    styles['CustomBodyText']
                ))
            elif '. ' in desc:
                for sent in [s.strip() for s in desc.split('. ') if s.strip()]:
                    story.append(Paragraph(f"• {sent}", styles['CustomBulletText']))
            else:
                story.append(Paragraph(desc, styles['CustomBodyText']))
            
            story.append(Spacer(1, 0.15 * inch))

    # ----------------  EDUCATION  -----------------
    if resume_data.get('education'):
        story.append(Paragraph("EDUCATION", styles['CustomSectionHeading']))
        for edu in resume_data['education']:
            story.append(Paragraph(f"<b>{edu['degree']}</b>", styles['CustomJobTitle']))
            
            inst_yr = edu['institution']
            if edu.get('year'):
                inst_yr += f" • {edu['year']}"
            story.append(Paragraph(inst_yr, styles['CustomCompanyDuration']))
            
            if edu.get('details') and len(edu['details']) > 20:
                story.append(Paragraph(
                    edu['details'][:200],
                    styles['CustomBodyText']
                ))
            story.append(Spacer(1, 0.12 * inch))

    # ----------------  SKILLS  -----------------
    if resume_data.get('skills'):
        story.append(Paragraph("SKILLS", styles['CustomSectionHeading']))
        story.append(Paragraph(
            " • ".join(resume_data['skills']),
            styles['CustomBodyText']
        ))

    doc.build(story)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = (
        f'attachment; filename="{user.username}_resume.pdf"'
    )
    return response
# ----------------------------------------------------------
#  DOCX export
# ----------------------------------------------------------
def generate_docx_resume(user, resume_data):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

   # --- header ---
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run('RESUME')
    run.font.size, run.font.bold, run.font.color.rgb = Pt(24), True, RGBColor(26, 37, 64)

    contact_parts = []
    for k in ('email', 'phone', 'location'):
        if resume_data.get(k):
            contact_parts.append(resume_data[k])
    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(' | '.join(contact_parts)).font.size = Pt(10)

    # --- sections ---
    def add_section(title):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size, run.font.bold, run.font.color.rgb = Pt(14), True, RGBColor(37, 99, 235)
        return p

    if resume_data.get('summary'):
        add_section('PROFESSIONAL SUMMARY')
        doc.add_paragraph(resume_data['summary']).paragraph_format.space_after = Pt(12)

    if resume_data.get('experience'):
        add_section('WORK EXPERIENCE')
        for exp in resume_data['experience']:
            doc.add_paragraph().add_run(exp['title']).font.bold = True
            company = exp['company'] + (f' • {exp["duration"]}' if exp.get('duration') else '')
            p = doc.add_paragraph(company); p.runs[0].font.italic = True
            desc = exp['description'].strip()
            if '. ' in desc and '•' not in desc:
                for sent in [s.strip() for s in desc.split('. ') if s.strip()]:
                    doc.add_paragraph(sent, style='List Bullet')
            else:
                doc.add_paragraph(desc)
            doc.add_paragraph()  # spacer

    if resume_data.get('education'):
        add_section('EDUCATION')
        for edu in resume_data['education']:
            doc.add_paragraph().add_run(edu['degree']).font.bold = True
            inst = edu['institution'] + (f' • {edu["year"]}' if edu.get('year') else '')
            doc.add_paragraph(inst).runs[0].font.italic = True
            doc.add_paragraph()

    if resume_data.get('skills'):
        add_section('SKILLS')
        doc.add_paragraph(' • '.join(resume_data['skills']))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = HttpResponse(buf.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp['Content-Disposition'] = f'attachment; filename="{user.username}_resume.docx"'
    return resp


# ----------------------------------------------------------
#  TXT export
# ----------------------------------------------------------
def generate_txt_resume(user, resume_data):
    lines = []
    def add(line=''): lines.append(line)

    add('=' * 70)
    add(resume_data['name'].upper().center(70))
    add('=' * 70); add()

    for k in ('email', 'phone', 'location'):
        if resume_data.get(k):
            add(f'{k.capitalize()}: {resume_data[k]}'.center(70))
    add(); add()

    if resume_data.get('summary'):
        add('PROFESSIONAL SUMMARY'); add('-' * 70)
        add(resume_data['summary']); add(); add()

    if resume_data.get('experience'):
        add('WORK EXPERIENCE'); add('-' * 70)
        for exp in resume_data['experience']:
            add(); add(exp['title'])
            company = exp['company'] + (f' | {exp["duration"]}' if exp.get('duration') else '')
            add(company); add()
            desc = exp['description'].strip()
            if '. ' in desc and '•' not in desc:
                for sent in [s.strip() for s in desc.split('. ') if s.strip()]:
                    add(f'  • {sent}')
            else:
                add('  ' + desc.replace('\n', '\n  '))
        add()

    if resume_data.get('education'):
        add('EDUCATION'); add('-' * 70)
        for edu in resume_data['education']:
            add(); add(edu['degree'])
            inst = edu['institution'] + (f' | {edu["year"]}' if edu.get('year') else '')
            add(inst); add()

    if resume_data.get('skills'):
        add('SKILLS'); add('-' * 70)
        add(' • '.join(resume_data['skills'])); add()

    add(); add('=' * 70)

    txt = '\n'.join(lines)
    resp = HttpResponse(txt, content_type='text/plain')
    resp['Content-Disposition'] = f'attachment; filename="{user.username}_resume.txt"'
    return resp


# ----------------------------------------------------------
#  Cover-letter helpers (PDF / DOCX / TXT)
# ----------------------------------------------------------
def generate_cover_letter_pdf(user, cover_data):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib import colors
    from io import BytesIO

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle('CL_Normal', parent=styles['Normal'], fontSize=11, leading=16, spaceAfter=12))
    styles.add(ParagraphStyle('CL_Right',  parent=styles['Normal'], fontSize=11, alignment=TA_RIGHT))

    story = []
    story.append(Paragraph(f"{cover_data['name']}<br/>{cover_data['email']}", styles['CL_Right']))
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph(cover_data['date'], styles['CL_Normal']))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(f"{cover_data['company_name']}<br/>{cover_data['position']}", styles['CL_Normal']))
    story.append(Spacer(1, 0.3 * inch))

    for para in cover_data['letter_body'].split('\n\n'):
        if para.strip():
            story.append(Paragraph(para.strip().replace('\n', '<br/>'), styles['CL_Normal']))
            story.append(Spacer(1, 0.15 * inch))

    doc.build(story)
    buf.seek(0)
    resp = HttpResponse(buf, content_type='application/pdf')
    resp['Content-Disposition'] = f'attachment; filename="{user.username}_cover_letter.pdf"'
    return resp


def generate_cover_letter_docx(user, cover_data):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"{cover_data['name']}\n{cover_data['email']}").font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph(cover_data['date'])
    doc.add_paragraph()
    doc.add_paragraph(cover_data['company_name'])
    doc.add_paragraph(cover_data['position'])
    doc.add_paragraph()

    for para in cover_data['letter_body'].split('\n\n'):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.space_after = Pt(12)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = HttpResponse(buf.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp['Content-Disposition'] = f'attachment; filename="{user.username}_cover_letter.docx"'
    return resp


def generate_cover_letter_txt(user, cover_data):
    lines = [
        cover_data['name'],
        cover_data['email'],
        '', '',
        cover_data['date'], '', '',
        cover_data['company_name'],
        cover_data['position'], '', '',
        cover_data['letter_body']
    ]
    txt = '\n'.join(lines)
    resp = HttpResponse(txt, content_type='text/plain')
    resp['Content-Disposition'] = f'attachment; filename="{user.username}_cover_letter.txt"'
    return resp

    
#export cover letter new
@login_required
def export_cover_letter(request):
    """
    Export cover letter as PDF, DOCX, or TXT.
    """
    format_type = request.GET.get("format", "pdf")
    
    # Get the cover letter
    cover_letter = Resume.objects.filter(
        user=request.user,
        status='draft',
        title__exact='Cover Letter Draft'
    ).first()

    if not cover_letter:
        return HttpResponse("No cover letter found. Please create one first.", status=404)

    # Extract cover letter data from conversation
    cover_data = extract_cover_letter_from_conversation(cover_letter)
    
    if not cover_data:
        return HttpResponse("No cover letter content found. Please complete the conversation.", status=404)

    # Generate the appropriate format
    if format_type == "pdf":
        return generate_cover_letter_pdf(request.user, cover_data)
    elif format_type == "docx":
        return generate_cover_letter_docx(request.user, cover_data)
    elif format_type == "txt":
        return generate_cover_letter_txt(request.user, cover_data)
    else:
        return HttpResponse("Invalid format", status=400)


def extract_cover_letter_from_conversation(cover_letter):
    """
    Extract structured cover letter data from conversation.
    """
    conversation = cover_letter.sections.get('conversation', [])
    
    if not conversation:
        return None
    
    # Get user responses
    user_responses = [msg['content'] for msg in conversation if msg.get('role') == 'user']
    ai_responses = [msg['content'] for msg in conversation if msg.get('role') == 'assistant']
    
    # Initialize cover letter data
    cover_data = {
        'name': cover_letter.user.get_full_name() or cover_letter.user.username,
        'email': cover_letter.user.email,
        'phone': '',
        'date': date.today().strftime("%B %d, %Y"),
        'company_name': 'Hiring Manager',
        'position': 'Position',
        'letter_body': ''
    }
    
    # Extract company and position from first user response
    if len(user_responses) > 0:
        first_response = user_responses[0].lower()
        
        # Try to extract company name
        if 'at' in first_response or 'with' in first_response:
            words = user_responses[0].split()
            for i, word in enumerate(words):
                if word.lower() in ['at', 'with'] and i + 1 < len(words):
                    cover_data['company_name'] = ' '.join(words[i+1:i+3])
                    break
        
        # Try to extract position
        for i, word in enumerate(user_responses[0].split()):
            if word.lower() in ['for', 'as']:
                cover_data['position'] = ' '.join(user_responses[0].split()[i+1:i+4])
                break
    
    # Get the final cover letter text
    # Look for the longest AI response (usually the final letter)
    final_letter = ""
    for response in reversed(ai_responses):
        if len(response) > 200:  # Substantial content
            lower = response.lower()
            # Check if it looks like a cover letter
            if any(phrase in lower for phrase in [
                'dear', 'sincerely', 'regards', 'position', 'experience', 'skills'
            ]):
                final_letter = response
                break
    
    # If no final letter found, combine last few AI responses
    if not final_letter and len(ai_responses) >= 2:
        final_letter = "\n\n".join(ai_responses[-2:])
    
    # If still nothing, create from user responses
    if not final_letter:
        final_letter = f"""Dear Hiring Manager,

I am writing to express my strong interest in the {cover_data['position']} position at {cover_data['company_name']}.

{' '.join(user_responses[:3])}

I am excited about the opportunity to contribute to your team and would welcome the chance to discuss how my skills align with your needs.

Thank you for considering my application.

Sincerely,
{cover_data['name']}"""
    
    cover_data['letter_body'] = final_letter
    
    return cover_data


def generate_cover_letter_pdf(user, cover_data):
    """
    Generate PDF cover letter.
    """

    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles for cover letter
    styles.add(ParagraphStyle(
        name='CoverLetterNormal',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        spaceAfter=12,
        alignment=TA_LEFT
    ))
    
    styles.add(ParagraphStyle(
        name='CoverLetterRight',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_RIGHT
    ))
    
    story = []
    
    # Your contact info (top right)
    contact_text = f"{cover_data['name']}<br/>{cover_data['email']}"
    if cover_data['phone']:
        contact_text += f"<br/>{cover_data['phone']}"
    story.append(Paragraph(contact_text, styles['CoverLetterRight']))
    story.append(Spacer(1, 0.3*inch))
    
    # Date
    story.append(Paragraph(cover_data['date'], styles['CoverLetterNormal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Recipient
    story.append(Paragraph(f"{cover_data['company_name']}<br/>{cover_data['position']}", styles['CoverLetterNormal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Letter body
    # Split into paragraphs
    paragraphs = cover_data['letter_body'].split('\n\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para.strip().replace('\n', '<br/>'), styles['CoverLetterNormal']))
            story.append(Spacer(1, 0.15*inch))
    
    doc.build(story)
    
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{user.username}_cover_letter.pdf"'
    return response


def generate_cover_letter_docx(user, cover_data):
    """
    Generate DOCX cover letter.
    """
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Your contact info (right aligned)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact_run = contact_para.add_run(f"{cover_data['name']}\n{cover_data['email']}")
    contact_run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacer
    
    # Date
    doc.add_paragraph(cover_data['date'])
    doc.add_paragraph()  # Spacer
    
    # Recipient
    doc.add_paragraph(cover_data['company_name'])
    doc.add_paragraph(cover_data['position'])
    doc.add_paragraph()  # Spacer
    
    # Letter body
    paragraphs = cover_data['letter_body'].split('\n\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.space_after = Pt(12)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{user.username}_cover_letter.docx"'
    return response


def generate_cover_letter_txt(user, cover_data):
    """
    Generate TXT cover letter.
    """
    lines = []
    
    # Your contact info
    lines.append(cover_data['name'])
    lines.append(cover_data['email'])
    if cover_data['phone']:
        lines.append(cover_data['phone'])
    lines.append("")
    lines.append("")
    
    # Date
    lines.append(cover_data['date'])
    lines.append("")
    lines.append("")
    
    # Recipient
    lines.append(cover_data['company_name'])
    lines.append(cover_data['position'])
    lines.append("")
    lines.append("")
    
    # Letter body
    lines.append(cover_data['letter_body'])
    
    text_content = "\n".join(lines)
    
    response = HttpResponse(text_content, content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename="{user.username}_cover_letter.txt"'
    return response