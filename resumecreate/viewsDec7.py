from django.shortcuts import render,redirect,get_object_or_404
from .models import Resume, UserProfile
from django.contrib.auth.models import User
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login, logout
from django.contrib.auth.views import LoginView, LogoutView
from django.contrib.auth.decorators import login_required
from django.urls import reverse_lazy
from django.shortcuts import redirect

from .models import UserProfile
from django import forms 


from django.http import HttpResponse
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
import json
from .services import OpenRouterService  # Import the service created
from django.template.loader import render_to_string

from django.views.decorators.http import require_POST

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, ListFlowable


from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    PageBreak,
    ListFlowable,
    ListItem
)

from reportlab.lib.pagesizes import letter
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.units import inch

from django.template.loader import render_to_string
from playwright.sync_api import sync_playwright
from datetime import date


# LOGIN VIEW
class CustomLoginView(LoginView):
    template_name = 'frontend/login.html'
    redirect_authenticated_user = True
    
    def get_success_url(self):
        return reverse_lazy('resumecreate:generate_cv')
    
    
# LOGOUT VIEW - ALTERNATIVE (Method 2 - Function-based - MORE RELIABLE)
def custom_logout(request):
    """
    Custom logout view that properly logs out and redirects.
    This is more reliable than the class-based view.
    """
    logout(request)
    return redirect('resumecreate:index')
    
    
# HOME VIEW
def home(request):
    """Render base.html as home/landing page"""
    context = {}
    return render(request, 'frontend/base.html', context)

# Admin dashboard view (renders base with all partials)
def admin_dashboard(request):
    registered_users = User.objects.select_related('profile').all()[:10]  # Joins profile
    submitted_resumes = Resume.objects.filter(status='submitted').order_by('-created_at')[:10]
    
    context = {
        'registered_users': registered_users, 
        'submitted_resumes': submitted_resumes,  
        'active_page': 'dashboard',
    }
    return render(request, 'admin/base_admin.html', context)

# View for registered users page (includes sidebar and user partial)
def admin_registered_users(request):
    registered_users = User.objects.select_related('profile').order_by('-date_joined')
    context = {
        'registered_users': registered_users,
        'active_page': 'users',  # Highlight in sidebar
    }
    return render(request, 'admin/base_admin.html', context)

# View for submitted resumes page
def admin_submitted_resumes(request):
    submitted_resumes = Resume.objects.filter(status='submitted').order_by('-created_at')
    context = {
        'submitted_resumes': submitted_resumes,
        'active_page': 'resumes',
    }
    return render(request, 'admin/base_admin.html', context)

# Detail view for a specific resume 
def admin_resume_detail(request, resume_id):
    # Find dummy resume or 404
    resume = next((r for r in DUMMY_RESUMES if r['id'] == resume_id), None)
    if not resume:
        # In real app, use get_object_or_404; here, dummy 404 context
        context = {'error': 'Resume not found'}
        return render(request, 'admin/partials/_resume_detail.html', context)
    
    context = {
        'resume': resume,
        'active_page': 'resumes',
    }
    return render(request, 'admin/partials/_resume_detail.html', context)  # Adjust template if needed

#E##########
class CustomUserCreationForm(UserCreationForm):
    email = forms.EmailField(required=True, label="Email Address")  # Add email field

    class Meta:
        model = User  # Django's User model
        fields = ('username', 'email', 'password1', 'password2')  # Include email

    def save(self, commit=True):
        user = super().save(commit=False)
        user.email = self.cleaned_data['email']
        if commit:
            user.save()
        return user

def register(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            UserProfile.objects.get_or_create(user=user)  # Auto-profile
            login(request, user)
            return redirect('resumecreate:generate_cv')  # Or your builder URL
    else:
        form = CustomUserCreationForm()
    return render(request, 'frontend/signup.html', {'form': form})


# def generate_cv(request): 
#     return render(request, 'frontend/generate_cv.html', {}) 


#reset conversation for resume builder
@login_required
@require_POST
def reset_resume_conversation(request):
    """
    Resets the current AI resume conversation and draft resume sections.
    """
    service = OpenRouterService()
    success = service.reset_user_conversation(request.user, request=request)

    if success:
        return JsonResponse({'status': 'success', 'message': 'Conversation reset'})
    else:
        return JsonResponse({'status': 'error', 'message': 'Failed to reset conversation'})





#generate cv view
@login_required
def generate_cv(request):
    """
    Main CV generation view with AI chat integration.
    COMPLETELY SEPARATE from cover letter builder.
    """
    user_initials = request.user.username[:2].upper()
    
    # IMPORTANT: Exclude cover letters from resume query
    resume = Resume.objects.filter(
        user=request.user, 
        status='draft'
    ).exclude(
        title__exact='Cover Letter Draft'  # Exclude cover letters
    ).order_by('-created_at').first()
    
    # Create new resume if none exists
    if not resume:
        resume = Resume.objects.create(
            user=request.user, 
            status='draft', 
            title='New Resume',
            sections={
                'type': 'resume',  # Mark as resume
                'conversation': [],
                'progress': {'current_step': 1}
            }
        )
    
    # Ensure sections has required structure
    if 'type' not in resume.sections:
        resume.sections['type'] = 'resume'
    if 'conversation' not in resume.sections:
        resume.sections['conversation'] = []
    if 'progress' not in resume.sections:
        resume.sections['progress'] = {'current_step': 1}
        resume.save()
    
    # Handle reset/clear conversation
    if request.GET.get('reset') == 'true' or request.POST.get('reset') == 'true':
        resume.sections = {
            'type': 'resume',
            'conversation': [],
            'progress': {'current_step': 1}
        }
        resume.save()
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': True, 'message': 'Conversation reset!'})
        return redirect('resumecreate:generate_cv')

    conversation_history = resume.sections.get('conversation', [])

    # Initialize progress data
    progress_data = {
        'current_step': resume.sections.get('progress', {}).get('current_step', 1),
        'total_steps': 6,
        'steps': [
            {"id": 1, "title": "Personal Info", "description": "Tell me your name, contact info, and location"},
            {"id": 2, "title": "Professional Summary", "description": "What's your career objective or professional summary?"},
            {"id": 3, "title": "Work Experience", "description": "Let's detail your work history and achievements"},
            {"id": 4, "title": "Education", "description": "Share your educational background"},
            {"id": 5, "title": "Skills & Certifications", "description": "What are your key skills and certifications?"},
            {"id": 6, "title": "Final Review", "description": "Review and finalize your resume"}
        ],
        'percentage': 16
    }

    if conversation_history:
        user_messages = [msg for msg in conversation_history if msg.get('role') == 'user']
        step_count = len(user_messages)
        current_step = min(step_count, 6) if step_count > 0 else 1
        progress_data['current_step'] = current_step
        progress_data['percentage'] = int((current_step / 6) * 100)
        resume.sections.setdefault('progress', {})['current_step'] = current_step
        resume.save()

    templates = [
        {'id': 1, 'image': 'frontend/images/1.png', 'name': 'Modern'},
        {'id': 2, 'image': 'frontend/images/2.png', 'name': 'Classic'},
        {'id': 3, 'image': 'frontend/images/3.png', 'name': 'Creative'},
        {'id': 4, 'image': 'frontend/images/4.png', 'name': 'Minimal'},
        {'id': 5, 'image': 'frontend/images/5.png', 'name': 'Professional'},
        {'id': 6, 'image': 'frontend/images/6.png', 'name': 'Tech'},
    ]

    ai_messages = []
    if not conversation_history or len(conversation_history) == 0:
        ai_messages = [{
            'avatar': 'AI',
            'text': 'Nice choice! Would you like to improve your current resume, or start from scratch?',
            'is_user': False,
            'show_actions': True
        }]
    else:
        for msg in conversation_history:
            if msg['role'] != 'system':
                ai_messages.append({
                    'avatar': user_initials if msg['role'] == 'user' else 'AI',
                    'text': msg['content'],
                    'is_user': msg['role'] == 'user',
                    'show_actions': False
                })

    if request.method == 'POST':
        try:
            ai_service = OpenRouterService()
            user_context = {
                'username': request.user.username,
                'email': request.user.email,
                'profile': {
                    'phone': getattr(request.user.profile, 'phone', ''),
                    'country': getattr(request.user.profile, 'country', ''),
                    'bio': getattr(request.user.profile, 'bio', '')
                } if hasattr(request.user, 'profile') else {},
                'current_step': progress_data['current_step'],
                'current_step_title': progress_data['steps'][progress_data['current_step']-1]['title']
            }

            # Handle template selection
            if 'template_id' in request.POST:
                template_id = request.POST.get('template_id')
                resume.sections['template_id'] = template_id
                resume.sections.setdefault('conversation', []).append({
                    'role': 'system',
                    'content': f'User selected template {template_id}'
                })
                resume.sections['type'] = 'resume'  # Ensure type is set
                resume.save()
                return JsonResponse({'success': True, 'message': 'Template selected successfully'})

            # Handle action buttons
            if 'action' in request.POST:
                action = request.POST.get('action')
                if action == 'improve':
                    user_message = "I'd like to improve my existing resume"
                elif action == 'fresh':
                    user_message = "I want to start from scratch"
                    progress_data['current_step'] = 1
                    progress_data['percentage'] = 16
                    resume.sections['progress']['current_step'] = 1
                else:
                    user_message = action

                conversation_history.append({'role': 'user', 'content': user_message})

                current_section = progress_data['steps'][progress_data['current_step']-1]['title'].lower()
                ai_response = ai_service.generate_resume_section(
                    request.user,
                    section_type=current_section,
                    user_data={"user_input": user_message}
                )

                conversation_history.append({'role': 'assistant', 'content': ai_response})

                # Advance progress
                if action != 'fresh':
                    progress_data['current_step'] = min(progress_data['current_step'] + 1, 6)
                progress_data['percentage'] = int((progress_data['current_step'] / 6) * 100)

                # Save updates
                resume.sections['conversation'] = conversation_history
                resume.sections['progress']['current_step'] = progress_data['current_step']
                resume.sections['type'] = 'resume'  # Ensure type is set
                resume.save()

                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': True,
                        'response': ai_response,
                        'avatar': 'AI',
                        'progress_data': progress_data
                    })

            # Handle chat input
            if 'chat_input' in request.POST:
                user_input = request.POST.get('chat_input', '').strip()
                if user_input:
                    conversation_history.append({'role': 'user', 'content': user_input})
                    current_section = progress_data['steps'][progress_data['current_step']-1]['title'].lower()
                    ai_response = ai_service.generate_resume_section(
                        request.user,
                        section_type=current_section,
                        user_data={"user_input": user_input}
                    )

                    conversation_history.append({'role': 'assistant', 'content': ai_response})

                    # Advance step if not last
                    if progress_data['current_step'] < 6:
                        progress_data['current_step'] += 1
                        progress_data['percentage'] = int((progress_data['current_step'] / 6) * 100)

                    # Save conversation
                    resume.sections['conversation'] = conversation_history
                    resume.sections['progress']['current_step'] = progress_data['current_step']
                    resume.sections['type'] = 'resume'  # Ensure type is set
                    resume.save()

                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return JsonResponse({
                            'success': True,
                            'response': ai_response,
                            'avatar': 'AI',
                            'progress_data': progress_data
                        })
                    else:
                        return redirect('resumecreate:generate_cv')

            # Handle export request
            if 'export' in request.POST:
                export_type = request.POST.get('export')
                resume.status = 'generated'
                resume.save()
                return JsonResponse({'success': True, 'message': f'Export to {export_type} requested'})
                
        except Exception as e:
            print(f"Resume generation error: {e}")  # Debug logging
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'error': str(e)})

    context = {
        'user_initials': user_initials,
        'templates': templates,
        'ai_messages': ai_messages,
        'resume_id': resume.id,
        'has_conversation': len(conversation_history) > 0,
        'current_step': progress_data['current_step'],
        'total_steps': progress_data['total_steps'],
        'steps': progress_data['steps'],
        'progress_percentage': progress_data['percentage']
    }

    return render(request, 'frontend/generate_cv.html', context)



#reset conversation
@login_required
def reset_conversation(request):
    """
    Resets the user's conversation and draft resume.
    """
    if request.method == 'POST':
        service = OpenRouterService()
        success = service.reset_user_conversation(request.user, request=request)
        return JsonResponse({'success': success})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=400)


#generate cover
@login_required
def generate_cover(request):
    """
    Cover letter generation view with AI integration and progress tracking.
    """
    user_initials = request.user.username[:2].upper()
    
    # Get or create draft cover letter (reusing Resume model)
    # Use a specific identifier to separate cover letters from resumes
    cover_letter = Resume.objects.filter(
        user=request.user,
        status='draft',
        title__exact='Cover Letter Draft'  # Exact match to avoid confusion
    ).order_by('-created_at').first()
    
    if not cover_letter:
        cover_letter = Resume.objects.create(
            user=request.user,
            status='draft',
            title='Cover Letter Draft',
            sections={'type': 'cover_letter', 'conversation': [], 'progress': {
                'job_details': False,
                'key_skills': False,
                'experience': False,
                'motivation': False
            }}
        )
    
    # Check if user wants to start fresh
    if request.GET.get('clear') == 'true':
        if cover_letter:
            cover_letter.sections = {
                'type': 'cover_letter',
                'conversation': [],
                'progress': {
                    'job_details': False,
                    'key_skills': False,
                    'experience': False,
                    'motivation': False
                }
            }
            cover_letter.save()
        return redirect('resumecreate:generate_cover')
    
    conversation_history = cover_letter.sections.get('conversation', [])
    
    # Calculate progress
    progress = cover_letter.sections.get('progress', {
        'job_details': False,
        'key_skills': False,
        'experience': False,
        'motivation': False
    })
    
    # Calculate completion percentage
    completed_items = sum(1 for v in progress.values() if v)
    completion_percentage = int((completed_items / 4) * 100)
    
    # Initial messages
    ai_messages = []
    if not conversation_history:
        ai_messages = [{
            'avatar': 'AI',
            'text': "Hello! I'm your AI cover letter assistant. To create a compelling cover letter, I'll need some information. Let's start with: What's the job title and company name you're applying to?",
            'is_user': False
        }]
    else:
        for msg in conversation_history:
            if msg['role'] != 'system':
                ai_messages.append({
                    'avatar': user_initials if msg['role'] == 'user' else 'AI',
                    'text': msg['content'],
                    'is_user': msg['role'] == 'user'
                })
    
    if request.method == 'POST':
        # Handle file upload
        if 'resume_file' in request.FILES:
            resume_file = request.FILES['resume_file']
            # TODO: Process resume file (extract text, save to model)
            # For now, just acknowledge receipt
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True,
                    'message': f'Resume "{resume_file.name}" uploaded successfully! Now, tell me about the job you\'re applying for.'
                })
        
        # Handle chat input
        user_input = request.POST.get('chat_input', '').strip()
        
        if user_input:
            conversation_history.append({'role': 'user', 'content': user_input})
            
            try:
                ai_service = OpenRouterService()
                
                # Cover letter specific system prompt
                system_prompt = """You are an expert cover letter writer. Help users create compelling, personalized cover letters.

Your approach:
1. Ask about job details (company, position, job description)
2. Inquire about relevant skills and experience
3. Understand their motivation for the role
4. Guide them to highlight achievements that match the job requirements

Keep responses conversational, encouraging, and professional. Ask one question at a time."""

                messages = [{'role': 'system', 'content': system_prompt}]
                
                # Add context if we have progress data
                if any(progress.values()):
                    context = f"Progress so far: {progress}"
                    messages.append({'role': 'system', 'content': context})
                
                messages.extend(conversation_history[:-1])
                messages.append({'role': 'user', 'content': user_input})
                
                ai_response = ai_service._make_request(messages)
                
                conversation_history.append({'role': 'assistant', 'content': ai_response})
                
                # Update progress based on conversation
                user_input_lower = user_input.lower()
                if any(word in user_input_lower for word in ['company', 'position', 'job', 'role', 'apply']):
                    progress['job_details'] = True
                if any(word in user_input_lower for word in ['skill', 'experience', 'python', 'javascript', 'manage', 'lead']):
                    progress['key_skills'] = True
                if any(word in user_input_lower for word in ['worked', 'years', 'developed', 'managed', 'led']):
                    progress['experience'] = True
                if any(word in user_input_lower for word in ['passionate', 'excited', 'because', 'want', 'interested']):
                    progress['motivation'] = True
                
                # Save to cover letter
                cover_letter.sections['conversation'] = conversation_history
                cover_letter.sections['progress'] = progress
                cover_letter.save()
                
                # Calculate new completion
                completed_items = sum(1 for v in progress.values() if v)
                completion_percentage = int((completed_items / 4) * 100)
                
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': True,
                        'response': ai_response,
                        'avatar': 'AI',
                        'progress': {
                            'job_details': progress['job_details'],
                            'key_skills': progress['key_skills'],
                            'experience': progress['experience'],
                            'motivation': progress['motivation'],
                            'percentage': completion_percentage
                        }
                    })
                else:
                    return redirect('resumecreate:generate_cover')
                    
            except Exception as e:
                error_message = f"Sorry, I encountered an error: {str(e)}"
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({'success': False, 'error': error_message})
        
        # Handle export
        if 'export' in request.POST:
            export_type = request.POST.get('export')
            cover_letter.status = 'generated'
            cover_letter.save()
            # TODO: Generate actual PDF/DOC file
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True, 
                    'message': f'Cover letter exported as {export_type}'
                })
    
    # Rebuild messages for display
    if conversation_history:
        ai_messages = []
        for msg in conversation_history:
            if msg['role'] != 'system':
                ai_messages.append({
                    'avatar': user_initials if msg['role'] == 'user' else 'AI',
                    'text': msg['content'],
                    'is_user': msg['role'] == 'user'
                })
    
    context = {
        'user_initials': user_initials,
        'ai_messages': ai_messages,
        'progress': progress,
        'completion_percentage': completion_percentage,
    }
    
    return render(request, 'frontend/generate_cover_letter.html', context)

    
#keep history for Gisele Knowles

@login_required
def user_history(request):
    user_initials = request.user.username[:2].upper() 

    resumes = Resume.objects.filter(user=request.user).order_by('-created_at')[:10]  

    context = {
        'user_initials': user_initials,
        'resumes': resumes,  # For card loop
    }

    if request.method == 'POST':
        search_query = request.POST.get('search')
        if search_query:
            resumes = resumes.filter(title__icontains=search_query)  # Filter by title
            context['resumes'] = resumes

    return render(request, 'frontend/history.html', context)


#resume detail

@login_required
def resume_detail(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)  # Secure: only own resumes
    context = {
        'resume': resume,  
    }
    return render(request, 'frontend/resume_detail.html', context)

#recommended Jobs
@login_required
def recommend_jobs_view(request):
    ai_service = OpenRouterService()
    user = request.user
    recommended_jobs = []
    current_job_title = ""

    if request.method == "POST":
        current_job_title = request.POST.get("current_job_title", "").strip()
        if current_job_title:
            # Get AI response string
            ai_response = ai_service.recommend_jobs(user, current_job_title)
            # Convert to a clean list
            recommended_jobs = [line.strip() for line in ai_response.split("\n") if line.strip()]

    return render(request, "frontend/recommended_jobs.html", {
        "recommended_jobs": recommended_jobs,
        "current_job_title": current_job_title,
        "user_initials": user.username[:2].upper(),
    })




#exporting resume to pdf, Beyonce
@login_required
def export_resume(request):
    """
    Export user's resume as a polished PDF or Word.
    """
    format = request.GET.get("format", "pdf")
    resume = Resume.objects.filter(user=request.user, status='draft').first()

    if not resume:
        return HttpResponse("No draft resume found.", status=404)

    if format == "pdf":
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=50, leftMargin=50,
                                topMargin=50, bottomMargin=50)

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='SectionHeading', fontSize=14, leading=16, spaceAfter=8, spaceBefore=12, fontName="Helvetica-Bold"))
        styles.add(ParagraphStyle(name='NormalText', fontSize=12, leading=14, spaceAfter=6))
        flowables = []

        # Title
        flowables.append(Paragraph(resume.title, styles['Title']))
        flowables.append(Spacer(1, 12))

        # Sections
        for section, content in resume.sections.items():
            flowables.append(Paragraph(section.title(), styles['SectionHeading']))

            if isinstance(content, str):
                # wrap long text in Paragraph
                flowables.append(Paragraph(content, styles['NormalText']))
            elif isinstance(content, (list, tuple)):
                # create bullets
                items = [ListItem(Paragraph(str(item), styles['NormalText']), bulletColor='black') for item in content]
                flowables.append(ListFlowable(items, bulletType='bullet', start='•', leftIndent=12))
            else:
                flowables.append(Paragraph(str(content), styles['NormalText']))

        doc.build(flowables)

        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="resume.pdf"'
        return response

    elif format == "doc":
        # Word export stays the same
        from docx import Document
        doc = Document()
        doc.add_heading(resume.title, level=0)
        for section, content in resume.sections.items():
            doc.add_heading(section.title(), level=1)
            if isinstance(content, str):
                doc.add_paragraph(content)
            elif isinstance(content, (list, tuple)):
                for item in content:
                    doc.add_paragraph(str(item), style='ListBullet')
            else:
                doc.add_paragraph(str(content))
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        response = HttpResponse(
            f.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="resume.docx"'
        return response

    else:
        return HttpResponse("Invalid format", status=400)
    
    
#export cover letter new
@login_required
def export_cover_letter_pdf(request):
    if request.method != "POST":
        return HttpResponse("Method not allowed", status=405)

    cover_letter_obj = Resume.objects.filter(
        user=request.user,
        status='draft',
        title='Cover Letter Draft'
    ).order_by('-created_at').first()

    if not cover_letter_obj:
        return HttpResponse("No cover letter found. Start a new one first.", status=404)

    conversation = cover_letter_obj.sections.get('conversation', [])

    # EXTRACT THE FINAL COVER LETTER — BULLETPROOF METHOD
    ai_responses = [m['content'].strip() for m in conversation if m.get('role') == 'assistant']
    
    final_letter = ""
    
    # Look for any message that contains the actual letter (most reliable)
    for response in reversed(ai_responses):
        lower = response.lower()
        if any(phrase in lower for phrase in [
            "here is your cover letter",
            "here's your cover letter",
            "final cover letter",
            "your cover letter",
            "dear hiring manager",
            "sincerely",
            "best regards"
        ]) or len(response) > 500:
            final_letter = response
            break

    # Fallback: combine last 3 AI messages (very common pattern)
    if not final_letter and len(ai_responses) >= 2:
        final_letter = "\n\n".join(ai_responses[-3:])

    # Final fallback
    if not final_letter:
        final_letter = "<p style='text-align:center; color:#666; padding:40px;'>Your cover letter is still being generated.<br><br>Keep chatting with the AI until it says 'Here is your final cover letter', then export.</p>"

    # Clean up markdown (optional but looks better)
    import re
    final_letter = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', final_letter)
    final_letter = re.sub(r'\*(.*?)\*', r'<em>\1</em>', final_letter)
    final_letter = final_letter.replace('\n', '<br>')

    context = {
        'cover_letter': final_letter,
        'user': request.user,
        'date': date.today().strftime("%B %d, %Y"),
        'company_name': 'Hiring Manager',
    }

    html_string = render_to_string('pdf/cover_letter_template.html', context)

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.set_content(html_string, wait_until="networkidle")
            pdf_bytes = page.pdf(
                format='A4',
                margin={'top': '1in', 'bottom': '1in', 'left': '0.8in', 'right': '0.8in'},
                print_background=True,
            )
            browser.close()

        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Cover_Letter_{request.user.username}_{date.today():%Y-%m-%d}.pdf"'
        
        # Mark as complete
        cover_letter_obj.status = 'generated'
        cover_letter_obj.save()

        return response

    except Exception as e:
        return HttpResponse(f"""
            <h2>PDF Generation Failed</h2>
            <p>Playwright error: {str(e)}</p>
            <p>Did you run: <code>playwright install chromium --with-deps</code> ?</p>
        """, status=500)