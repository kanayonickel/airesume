from django.shortcuts import render,redirect,get_object_or_404
from .models import Resume, UserProfile
from django.contrib.auth.models import User
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login, logout
from django.contrib.auth.views import LoginView, LogoutView
from django.contrib.auth.decorators import login_required
from django.urls import reverse_lazy
from django.shortcuts import redirect

from .models import UserProfile
from django import forms 


from django.http import HttpResponse
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
import json
from .services import OpenRouterService  # Import the service created
from django.template.loader import render_to_string

from django.views.decorators.http import require_POST

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
# from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, ListFlowable


# from reportlab.platypus import (
#     SimpleDocTemplate,
#     Paragraph,
#     Spacer,
#     PageBreak,
#     ListFlowable,
#     ListItem
# )

from reportlab.lib.pagesizes import letter
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.units import inch

from django.template.loader import render_to_string
from playwright.sync_api import sync_playwright
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors


# LOGIN VIEW
class CustomLoginView(LoginView):
    template_name = 'frontend/login.html'
    redirect_authenticated_user = True
    
    def get_success_url(self):
        return reverse_lazy('resumecreate:generate_cv')
    
    
# LOGOUT VIEW - ALTERNATIVE (Method 2 - Function-based - MORE RELIABLE)
def custom_logout(request):
    """
    Custom logout view that properly logs out and redirects.
    This is more reliable than the class-based view.
    """
    logout(request)
    return redirect('resumecreate:index')
    
    
# HOME VIEW
def home(request):
    """Render base.html as home/landing page"""
    context = {}
    return render(request, 'frontend/base.html', context)

# Admin dashboard view (renders base with all partials)
def admin_dashboard(request):
    registered_users = User.objects.select_related('profile').all()[:10]  # Joins profile
    submitted_resumes = Resume.objects.filter(status='submitted').order_by('-created_at')[:10]
    
    context = {
        'registered_users': registered_users, 
        'submitted_resumes': submitted_resumes,  
        'active_page': 'dashboard',
    }
    return render(request, 'admin/base_admin.html', context)

# View for registered users page (includes sidebar and user partial)
def admin_registered_users(request):
    registered_users = User.objects.select_related('profile').order_by('-date_joined')
    context = {
        'registered_users': registered_users,
        'active_page': 'users',  # Highlight in sidebar
    }
    return render(request, 'admin/base_admin.html', context)

# View for submitted resumes page
def admin_submitted_resumes(request):
    submitted_resumes = Resume.objects.filter(status='submitted').order_by('-created_at')
    context = {
        'submitted_resumes': submitted_resumes,
        'active_page': 'resumes',
    }
    return render(request, 'admin/base_admin.html', context)

# Detail view for a specific resume 
def admin_resume_detail(request, resume_id):
    # Find dummy resume or 404
    resume = next((r for r in DUMMY_RESUMES if r['id'] == resume_id), None)
    if not resume:
        # In real app, use get_object_or_404; here, dummy 404 context
        context = {'error': 'Resume not found'}
        return render(request, 'admin/partials/_resume_detail.html', context)
    
    context = {
        'resume': resume,
        'active_page': 'resumes',
    }
    return render(request, 'admin/partials/_resume_detail.html', context)  # Adjust template if needed

#E##########
class CustomUserCreationForm(UserCreationForm):
    email = forms.EmailField(required=True, label="Email Address")  # Add email field

    class Meta:
        model = User  # Django's User model
        fields = ('username', 'email', 'password1', 'password2')  # Include email

    def save(self, commit=True):
        user = super().save(commit=False)
        user.email = self.cleaned_data['email']
        if commit:
            user.save()
        return user

def register(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            UserProfile.objects.get_or_create(user=user)  # Auto-profile
            login(request, user)
            return redirect('resumecreate:generate_cv')  # Or your builder URL
    else:
        form = CustomUserCreationForm()
    return render(request, 'frontend/signup.html', {'form': form})


# def generate_cv(request): 
#     return render(request, 'frontend/generate_cv.html', {}) 


#reset conversation for resume builder
@login_required
@require_POST
def reset_resume_conversation(request):
    """
    Resets the current AI resume conversation and draft resume sections.
    """
    service = OpenRouterService()
    success = service.reset_user_conversation(request.user, request=request)

    if success:
        return JsonResponse({'status': 'success', 'message': 'Conversation reset'})
    else:
        return JsonResponse({'status': 'error', 'message': 'Failed to reset conversation'})





#generate cv view
@login_required
def generate_cv(request):
    """
    Main CV generation view with AI chat integration.
    COMPLETELY SEPARATE from cover letter builder.
    """
    user_initials = request.user.username[:2].upper()
    
    # IMPORTANT: Exclude cover letters from resume query
    resume = Resume.objects.filter(
        user=request.user, 
        status='draft'
    ).exclude(
        title__exact='Cover Letter Draft'  # Exclude cover letters
    ).order_by('-created_at').first()
    
    # Create new resume if none exists
    if not resume:
        resume = Resume.objects.create(
            user=request.user, 
            status='draft', 
            title='New Resume',
            sections={
                'type': 'resume',  # Mark as resume
                'conversation': [],
                'progress': {'current_step': 1}
            }
        )
    
    # Ensure sections has required structure
    if 'type' not in resume.sections:
        resume.sections['type'] = 'resume'
    if 'conversation' not in resume.sections:
        resume.sections['conversation'] = []
    if 'progress' not in resume.sections:
        resume.sections['progress'] = {'current_step': 1}
        resume.save()
    
    # Handle reset/clear conversation
    if request.GET.get('reset') == 'true' or request.POST.get('reset') == 'true':
        resume.sections = {
            'type': 'resume',
            'conversation': [],
            'progress': {'current_step': 1}
        }
        resume.save()
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': True, 'message': 'Conversation reset!'})
        return redirect('resumecreate:generate_cv')

    conversation_history = resume.sections.get('conversation', [])

    # Initialize progress data
    progress_data = {
        'current_step': resume.sections.get('progress', {}).get('current_step', 1),
        'total_steps': 6,
        'steps': [
            {"id": 1, "title": "Personal Info", "description": "Tell me your name, contact info, and location"},
            {"id": 2, "title": "Professional Summary", "description": "What's your career objective or professional summary?"},
            {"id": 3, "title": "Work Experience", "description": "Let's detail your work history and achievements"},
            {"id": 4, "title": "Education", "description": "Share your educational background"},
            {"id": 5, "title": "Skills & Certifications", "description": "What are your key skills and certifications?"},
            {"id": 6, "title": "Final Review", "description": "Review and finalize your resume"}
        ],
        'percentage': 16
    }

    if conversation_history:
        user_messages = [msg for msg in conversation_history if msg.get('role') == 'user']
        step_count = len(user_messages)
        current_step = min(step_count, 6) if step_count > 0 else 1
        progress_data['current_step'] = current_step
        progress_data['percentage'] = int((current_step / 6) * 100)
        resume.sections.setdefault('progress', {})['current_step'] = current_step
        resume.save()

    templates = [
        {'id': 1, 'image': 'frontend/images/1.png', 'name': 'Modern'},
        {'id': 2, 'image': 'frontend/images/2.png', 'name': 'Classic'},
        {'id': 3, 'image': 'frontend/images/3.png', 'name': 'Creative'},
        {'id': 4, 'image': 'frontend/images/4.png', 'name': 'Minimal'},
        {'id': 5, 'image': 'frontend/images/5.png', 'name': 'Professional'},
        {'id': 6, 'image': 'frontend/images/6.png', 'name': 'Tech'},
    ]

    ai_messages = []
    if not conversation_history or len(conversation_history) == 0:
        ai_messages = [{
            'avatar': 'AI',
            'text': 'Nice choice! Would you like to improve your current resume, or start from scratch?',
            'is_user': False,
            'show_actions': True
        }]
    else:
        for msg in conversation_history:
            if msg['role'] != 'system':
                ai_messages.append({
                    'avatar': user_initials if msg['role'] == 'user' else 'AI',
                    'text': msg['content'],
                    'is_user': msg['role'] == 'user',
                    'show_actions': False
                })

    if request.method == 'POST':
        try:
            ai_service = OpenRouterService()
            user_context = {
                'username': request.user.username,
                'email': request.user.email,
                'profile': {
                    'phone': getattr(request.user.profile, 'phone', ''),
                    'country': getattr(request.user.profile, 'country', ''),
                    'bio': getattr(request.user.profile, 'bio', '')
                } if hasattr(request.user, 'profile') else {},
                'current_step': progress_data['current_step'],
                'current_step_title': progress_data['steps'][progress_data['current_step']-1]['title']
            }

            # Handle template selection
            if 'template_id' in request.POST:
                template_id = request.POST.get('template_id')
                resume.sections['template_id'] = template_id
                resume.sections.setdefault('conversation', []).append({
                    'role': 'system',
                    'content': f'User selected template {template_id}'
                })
                resume.sections['type'] = 'resume'  # Ensure type is set
                resume.save()
                return JsonResponse({'success': True, 'message': 'Template selected successfully'})

            # Handle action buttons
            if 'action' in request.POST:
                action = request.POST.get('action')
                if action == 'improve':
                    user_message = "I'd like to improve my existing resume"
                elif action == 'fresh':
                    user_message = "I want to start from scratch"
                    progress_data['current_step'] = 1
                    progress_data['percentage'] = 16
                    resume.sections['progress']['current_step'] = 1
                else:
                    user_message = action

                conversation_history.append({'role': 'user', 'content': user_message})

                current_section = progress_data['steps'][progress_data['current_step']-1]['title'].lower()
                ai_response = ai_service.generate_resume_section(
                    request.user,
                    section_type=current_section,
                    user_data={"user_input": user_message}
                )

                conversation_history.append({'role': 'assistant', 'content': ai_response})

                # Advance progress
                if action != 'fresh':
                    progress_data['current_step'] = min(progress_data['current_step'] + 1, 6)
                progress_data['percentage'] = int((progress_data['current_step'] / 6) * 100)

                # Save updates
                resume.sections['conversation'] = conversation_history
                resume.sections['progress']['current_step'] = progress_data['current_step']
                resume.sections['type'] = 'resume'  # Ensure type is set
                resume.save()

                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': True,
                        'response': ai_response,
                        'avatar': 'AI',
                        'progress_data': progress_data
                    })

            # Handle chat input
            if 'chat_input' in request.POST:
                user_input = request.POST.get('chat_input', '').strip()
                if user_input:
                    conversation_history.append({'role': 'user', 'content': user_input})
                    current_section = progress_data['steps'][progress_data['current_step']-1]['title'].lower()
                    ai_response = ai_service.generate_resume_section(
                        request.user,
                        section_type=current_section,
                        user_data={"user_input": user_input}
                    )

                    conversation_history.append({'role': 'assistant', 'content': ai_response})

                    # Advance step if not last
                    if progress_data['current_step'] < 6:
                        progress_data['current_step'] += 1
                        progress_data['percentage'] = int((progress_data['current_step'] / 6) * 100)

                    # Save conversation
                    resume.sections['conversation'] = conversation_history
                    resume.sections['progress']['current_step'] = progress_data['current_step']
                    resume.sections['type'] = 'resume'  # Ensure type is set
                    resume.save()

                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return JsonResponse({
                            'success': True,
                            'response': ai_response,
                            'avatar': 'AI',
                            'progress_data': progress_data
                        })
                    else:
                        return redirect('resumecreate:generate_cv')

            # Handle export request
            if 'export' in request.POST:
                export_type = request.POST.get('export')
                resume.status = 'generated'
                resume.save()
                return JsonResponse({'success': True, 'message': f'Export to {export_type} requested'})
                
        except Exception as e:
            print(f"Resume generation error: {e}")  # Debug logging
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'error': str(e)})

    context = {
        'user_initials': user_initials,
        'templates': templates,
        'ai_messages': ai_messages,
        'resume_id': resume.id,
        'has_conversation': len(conversation_history) > 0,
        'current_step': progress_data['current_step'],
        'total_steps': progress_data['total_steps'],
        'steps': progress_data['steps'],
        'progress_percentage': progress_data['percentage']
    }

    return render(request, 'frontend/generate_cv.html', context)



#reset conversation
@login_required
def reset_conversation(request):
    """
    Resets the user's conversation and draft resume.
    """
    if request.method == 'POST':
        service = OpenRouterService()
        success = service.reset_user_conversation(request.user, request=request)
        return JsonResponse({'success': success})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=400)


#generate cover
@login_required
def generate_cover(request):
    """
    Cover letter generation view with AI integration and progress tracking.
    """
    user_initials = request.user.username[:2].upper()
    
    # Get or create draft cover letter (reusing Resume model)
    # Use a specific identifier to separate cover letters from resumes
    cover_letter = Resume.objects.filter(
        user=request.user,
        status='draft',
        title__exact='Cover Letter Draft'  # Exact match to avoid confusion
    ).order_by('-created_at').first()
    
    if not cover_letter:
        cover_letter = Resume.objects.create(
            user=request.user,
            status='draft',
            title='Cover Letter Draft',
            sections={'type': 'cover_letter', 'conversation': [], 'progress': {
                'job_details': False,
                'key_skills': False,
                'experience': False,
                'motivation': False
            }}
        )
    
    # Check if user wants to start fresh
    if request.GET.get('clear') == 'true':
        if cover_letter:
            cover_letter.sections = {
                'type': 'cover_letter',
                'conversation': [],
                'progress': {
                    'job_details': False,
                    'key_skills': False,
                    'experience': False,
                    'motivation': False
                }
            }
            cover_letter.save()
        return redirect('resumecreate:generate_cover')
    
    conversation_history = cover_letter.sections.get('conversation', [])
    
    # Calculate progress
    progress = cover_letter.sections.get('progress', {
        'job_details': False,
        'key_skills': False,
        'experience': False,
        'motivation': False
    })
    
    # Calculate completion percentage
    completed_items = sum(1 for v in progress.values() if v)
    completion_percentage = int((completed_items / 4) * 100)
    
    # Initial messages
    ai_messages = []
    if not conversation_history:
        ai_messages = [{
            'avatar': 'AI',
            'text': "Hello! I'm your AI cover letter assistant. To create a compelling cover letter, I'll need some information. Let's start with: What's the job title and company name you're applying to?",
            'is_user': False
        }]
    else:
        for msg in conversation_history:
            if msg['role'] != 'system':
                ai_messages.append({
                    'avatar': user_initials if msg['role'] == 'user' else 'AI',
                    'text': msg['content'],
                    'is_user': msg['role'] == 'user'
                })
    
    if request.method == 'POST':
        # Handle file upload
        if 'resume_file' in request.FILES:
            resume_file = request.FILES['resume_file']
            # TODO: Process resume file (extract text, save to model)
            # For now, just acknowledge receipt
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True,
                    'message': f'Resume "{resume_file.name}" uploaded successfully! Now, tell me about the job you\'re applying for.'
                })
        
        # Handle chat input
        user_input = request.POST.get('chat_input', '').strip()
        
        if user_input:
            conversation_history.append({'role': 'user', 'content': user_input})
            
            try:
                ai_service = OpenRouterService()
                
                # Cover letter specific system prompt
                system_prompt = """You are an expert cover letter writer. Help users create compelling, personalized cover letters.

Your approach:
1. Ask about job details (company, position, job description)
2. Inquire about relevant skills and experience
3. Understand their motivation for the role
4. Guide them to highlight achievements that match the job requirements

Keep responses conversational, encouraging, and professional. Ask one question at a time."""

                messages = [{'role': 'system', 'content': system_prompt}]
                
                # Add context if we have progress data
                if any(progress.values()):
                    context = f"Progress so far: {progress}"
                    messages.append({'role': 'system', 'content': context})
                
                messages.extend(conversation_history[:-1])
                messages.append({'role': 'user', 'content': user_input})
                
                ai_response = ai_service._make_request(messages)
                
                conversation_history.append({'role': 'assistant', 'content': ai_response})
                
                # Update progress based on conversation
                user_input_lower = user_input.lower()
                if any(word in user_input_lower for word in ['company', 'position', 'job', 'role', 'apply']):
                    progress['job_details'] = True
                if any(word in user_input_lower for word in ['skill', 'experience', 'python', 'javascript', 'manage', 'lead']):
                    progress['key_skills'] = True
                if any(word in user_input_lower for word in ['worked', 'years', 'developed', 'managed', 'led']):
                    progress['experience'] = True
                if any(word in user_input_lower for word in ['passionate', 'excited', 'because', 'want', 'interested']):
                    progress['motivation'] = True
                
                # Save to cover letter
                cover_letter.sections['conversation'] = conversation_history
                cover_letter.sections['progress'] = progress
                cover_letter.save()
                
                # Calculate new completion
                completed_items = sum(1 for v in progress.values() if v)
                completion_percentage = int((completed_items / 4) * 100)
                
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': True,
                        'response': ai_response,
                        'avatar': 'AI',
                        'progress': {
                            'job_details': progress['job_details'],
                            'key_skills': progress['key_skills'],
                            'experience': progress['experience'],
                            'motivation': progress['motivation'],
                            'percentage': completion_percentage
                        }
                    })
                else:
                    return redirect('resumecreate:generate_cover')
                    
            except Exception as e:
                error_message = f"Sorry, I encountered an error: {str(e)}"
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({'success': False, 'error': error_message})
        
        # Handle export
        if 'export' in request.POST:
            export_type = request.POST.get('export')
            cover_letter.status = 'generated'
            cover_letter.save()
            # TODO: Generate actual PDF/DOC file
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True, 
                    'message': f'Cover letter exported as {export_type}'
                })
    
    # Rebuild messages for display
    if conversation_history:
        ai_messages = []
        for msg in conversation_history:
            if msg['role'] != 'system':
                ai_messages.append({
                    'avatar': user_initials if msg['role'] == 'user' else 'AI',
                    'text': msg['content'],
                    'is_user': msg['role'] == 'user'
                })
    
    context = {
        'user_initials': user_initials,
        'ai_messages': ai_messages,
        'progress': progress,
        'completion_percentage': completion_percentage,
    }
    
    return render(request, 'frontend/generate_cover_letter.html', context)

    
#keep history for Gisele Knowles

@login_required
def user_history(request):
    user_initials = request.user.username[:2].upper() 

    resumes = Resume.objects.filter(user=request.user).order_by('-created_at')[:10]  

    context = {
        'user_initials': user_initials,
        'resumes': resumes,  # For card loop
    }

    if request.method == 'POST':
        search_query = request.POST.get('search')
        if search_query:
            resumes = resumes.filter(title__icontains=search_query)  # Filter by title
            context['resumes'] = resumes

    return render(request, 'frontend/history.html', context)


#resume detail

@login_required
def resume_detail(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)  # Secure: only own resumes
    context = {
        'resume': resume,  
    }
    return render(request, 'frontend/resume_detail.html', context)

#recommended Jobs
@login_required
def recommend_jobs_view(request):
    ai_service = OpenRouterService()
    user = request.user
    recommended_jobs = []
    current_job_title = ""

    if request.method == "POST":
        current_job_title = request.POST.get("current_job_title", "").strip()
        if current_job_title:
            # Get AI response string
            ai_response = ai_service.recommend_jobs(user, current_job_title)
            # Convert to a clean list
            recommended_jobs = [line.strip() for line in ai_response.split("\n") if line.strip()]

    return render(request, "frontend/recommended_jobs.html", {
        "recommended_jobs": recommended_jobs,
        "current_job_title": current_job_title,
        "user_initials": user.username[:2].upper(),
    })




#exporting resume to pdf, Beyonce
@login_required
def export_resume(request):
    """
    Export user's AI-generated resume as PDF, DOCX, or TXT.
    Extracts content from conversation history.
    """
    format_type = request.GET.get("format", "pdf")
    
    # Get the resume
    resume = Resume.objects.filter(
        user=request.user, 
        status='draft'
    ).exclude(
        title__exact='Cover Letter Draft'
    ).order_by('-created_at').first()

    if not resume:
        return HttpResponse("No draft resume found. Please create a resume first.", status=404)

    # Extract resume data from conversation
    resume_data = extract_resume_from_conversation(resume)
    
    if not resume_data:
        return HttpResponse("No resume content found. Please complete the resume building process.", status=404)

    # Generate the appropriate format
    if format_type == "pdf":
        return generate_pdf_resume(request.user, resume_data)
    elif format_type == "doc":
        return generate_docx_resume(request.user, resume_data)
    elif format_type == "txt":
        return generate_txt_resume(request.user, resume_data)
    else:
        return HttpResponse("Invalid format", status=400)


def extract_resume_from_conversation(resume):
    """
    Extract structured resume data from AI conversation history.
    Returns a dict with resume sections.
    """
    conversation = resume.sections.get('conversation', [])
    
    if not conversation:
        return None
    
    # Get all AI responses
    ai_responses = [msg['content'] for msg in conversation if msg.get('role') == 'assistant']
    
    # Combine all AI responses
    full_text = "\n\n".join(ai_responses)
    
    # Initialize resume data structure
    resume_data = {
        'name': '',
        'email': '',
        'phone': '',
        'location': '',
        'linkedin': '',
        'summary': '',
        'experience': [],
        'education': [],
        'skills': [],
        'certifications': []
    }
    
    # Extract user messages for context
    user_responses = [msg['content'] for msg in conversation if msg.get('role') == 'user']
    
    # Smart extraction from conversation
    # 1. Extract contact info from early messages
    if len(user_responses) > 0:
        first_response = user_responses[0]
        
        # Extract email
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', first_response)
        if email_match:
            resume_data['email'] = email_match.group()
        
        # Extract phone
        phone_match = re.search(r'\b\d{10,11}\b|\b\(\d{3}\)\s*\d{3}-\d{4}\b', first_response)
        if phone_match:
            resume_data['phone'] = phone_match.group()
        
        # Try to get name (usually first few words)
        words = first_response.split()
        if len(words) >= 2:
            potential_name = ' '.join(words[:2])
            if not any(char.isdigit() for char in potential_name) and '@' not in potential_name:
                resume_data['name'] = potential_name
    
    # If name not found, use username
    if not resume_data['name']:
        resume_data['name'] = resume.user.get_full_name() or resume.user.username
    
    # If email not found, use user email
    if not resume_data['email']:
        resume_data['email'] = resume.user.email
    
    # 2. Extract summary (usually in step 2 responses)
    for response in ai_responses:
        if len(response) > 50 and len(response) < 500 and ('summary' in response.lower() or 'objective' in response.lower()):
            resume_data['summary'] = response.strip()
            break
    
    # If no explicit summary found, create one from user responses
    if not resume_data['summary'] and len(user_responses) > 1:
        resume_data['summary'] = f"Professional with experience in {user_responses[1][:100]}..."
    
    # 3. Extract experience (look for job-related keywords)
    experience_items = []
    for i, response in enumerate(user_responses):
        lower = response.lower()
        if any(word in lower for word in ['worked', 'developer', 'engineer', 'manager', 'analyst', 'consultant', 'years']):
            # This might be experience info
            experience_items.append(response)
    
    # Parse experience items
    for item in experience_items:
        exp = {
            'title': 'Professional Role',
            'company': 'Company Name',
            'duration': '2020 - Present',
            'description': item[:200]
        }
        
        # Try to extract job title
        title_match = re.search(r'(software developer|engineer|manager|analyst|consultant|designer)', item, re.IGNORECASE)
        if title_match:
            exp['title'] = title_match.group().title()
        
        resume_data['experience'].append(exp)
    
    # 4. Extract education
    education_keywords = ['degree', 'bachelor', 'master', 'phd', 'university', 'college', 'graduated']
    for response in user_responses:
        if any(word in response.lower() for word in education_keywords):
            edu = {
                'degree': 'Bachelor of Science',
                'institution': 'University',
                'year': '2020',
                'details': response[:150]
            }
            resume_data['education'].append(edu)
    
    # 5. Extract skills
    skill_keywords = ['python', 'javascript', 'java', 'react', 'django', 'sql', 'aws', 'docker', 
                      'leadership', 'communication', 'teamwork', 'problem solving']
    
    all_text = ' '.join(user_responses).lower()
    found_skills = [skill for skill in skill_keywords if skill in all_text]
    
    if found_skills:
        resume_data['skills'] = [skill.title() for skill in found_skills]
    else:
        # Extract from last few responses
        for response in user_responses[-3:]:
            words = response.split(',')
            for word in words:
                clean = word.strip().title()
                if len(clean) > 2 and len(clean) < 30:
                    resume_data['skills'].append(clean)
    
    # Limit skills to reasonable number
    resume_data['skills'] = resume_data['skills'][:15]
    
    return resume_data


def generate_pdf_resume(user, resume_data):
    """
    Generate a PDF resume using ReportLab.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter,
        rightMargin=0.75*inch, 
        leftMargin=0.75*inch,
        topMargin=0.75*inch, 
        bottomMargin=0.75*inch
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    styles.add(ParagraphStyle(
        name='CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a2540'),
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='Contact',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=12
    ))
    
    styles.add(ParagraphStyle(
        name='SectionHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#2563EB'),
        spaceAfter=8,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    ))
    
    story = []
    
    # Header - Name
    story.append(Paragraph(resume_data['name'], styles['CustomTitle']))
    
    # Contact Info
    contact_parts = []
    if resume_data['email']:
        contact_parts.append(resume_data['email'])
    if resume_data['phone']:
        contact_parts.append(resume_data['phone'])
    if resume_data['location']:
        contact_parts.append(resume_data['location'])
    
    contact_text = " | ".join(contact_parts)
    story.append(Paragraph(contact_text, styles['Contact']))
    story.append(Spacer(1, 0.2*inch))
    
    # Summary
    if resume_data['summary']:
        story.append(Paragraph("PROFESSIONAL SUMMARY", styles['SectionHeading']))
        story.append(Paragraph(resume_data['summary'], styles['Normal']))
        story.append(Spacer(1, 0.15*inch))
    
    # Experience
    if resume_data['experience']:
        story.append(Paragraph("WORK EXPERIENCE", styles['SectionHeading']))
        for exp in resume_data['experience']:
            # Job title and company
            title_text = f"<b>{exp['title']}</b> - {exp['company']}"
            story.append(Paragraph(title_text, styles['Normal']))
            # Duration
            story.append(Paragraph(f"<i>{exp['duration']}</i>", styles['Normal']))
            # Description
            story.append(Paragraph(exp['description'], styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
    
    # Education
    if resume_data['education']:
        story.append(Paragraph("EDUCATION", styles['SectionHeading']))
        for edu in resume_data['education']:
            edu_text = f"<b>{edu['degree']}</b> - {edu['institution']} ({edu['year']})"
            story.append(Paragraph(edu_text, styles['Normal']))
            if edu.get('details'):
                story.append(Paragraph(edu['details'], styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
    
    # Skills
    if resume_data['skills']:
        story.append(Paragraph("SKILLS", styles['SectionHeading']))
        skills_text = " • ".join(resume_data['skills'])
        story.append(Paragraph(skills_text, styles['Normal']))
    
    # Build PDF
    doc.build(story)
    
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{user.username}_resume.pdf"'
    return response


def generate_docx_resume(user, resume_data):
    """
    Generate a DOCX resume using python-docx.
    """
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Name (centered, large, bold)
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(resume_data['name'])
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(26, 37, 64)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info (centered)
    contact_parts = []
    if resume_data['email']:
        contact_parts.append(resume_data['email'])
    if resume_data['phone']:
        contact_parts.append(resume_data['phone'])
    if resume_data['location']:
        contact_parts.append(resume_data['location'])
    
    contact_para = doc.add_paragraph(' | '.join(contact_parts))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # Summary
    if resume_data['summary']:
        heading = doc.add_heading('PROFESSIONAL SUMMARY', level=2)
        heading.runs[0].font.color.rgb = RGBColor(37, 99, 235)
        doc.add_paragraph(resume_data['summary'])
    
    # Experience
    if resume_data['experience']:
        heading = doc.add_heading('WORK EXPERIENCE', level=2)
        heading.runs[0].font.color.rgb = RGBColor(37, 99, 235)
        
        for exp in resume_data['experience']:
            # Job title (bold)
            job_para = doc.add_paragraph()
            job_run = job_para.add_run(f"{exp['title']} - {exp['company']}")
            job_run.bold = True
            
            # Duration (italic)
            dur_para = doc.add_paragraph()
            dur_run = dur_para.add_run(exp['duration'])
            dur_run.italic = True
            
            # Description
            doc.add_paragraph(exp['description'])
            doc.add_paragraph()  # Spacer
    
    # Education
    if resume_data['education']:
        heading = doc.add_heading('EDUCATION', level=2)
        heading.runs[0].font.color.rgb = RGBColor(37, 99, 235)
        
        for edu in resume_data['education']:
            edu_para = doc.add_paragraph()
            edu_run = edu_para.add_run(f"{edu['degree']} - {edu['institution']} ({edu['year']})")
            edu_run.bold = True
            
            if edu.get('details'):
                doc.add_paragraph(edu['details'])
            doc.add_paragraph()  # Spacer
    
    # Skills
    if resume_data['skills']:
        heading = doc.add_heading('SKILLS', level=2)
        heading.runs[0].font.color.rgb = RGBColor(37, 99, 235)
        doc.add_paragraph(' • '.join(resume_data['skills']))
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{user.username}_resume.docx"'
    return response


def generate_txt_resume(user, resume_data):
    """
    Generate a plain text resume.
    """
    lines = []
    
    # Name
    lines.append("=" * 70)
    lines.append(resume_data['name'].upper().center(70))
    lines.append("=" * 70)
    lines.append("")
    
    # Contact
    contact_parts = []
    if resume_data['email']:
        contact_parts.append(f"Email: {resume_data['email']}")
    if resume_data['phone']:
        contact_parts.append(f"Phone: {resume_data['phone']}")
    if resume_data['location']:
        contact_parts.append(f"Location: {resume_data['location']}")
    
    lines.append(" | ".join(contact_parts).center(70))
    lines.append("")
    lines.append("")
    
    # Summary
    if resume_data['summary']:
        lines.append("PROFESSIONAL SUMMARY")
        lines.append("-" * 70)
        lines.append(resume_data['summary'])
        lines.append("")
        lines.append("")
    
    # Experience
    if resume_data['experience']:
        lines.append("WORK EXPERIENCE")
        lines.append("-" * 70)
        for exp in resume_data['experience']:
            lines.append(f"{exp['title']} - {exp['company']}")
            lines.append(exp['duration'])
            lines.append(exp['description'])
            lines.append("")
        lines.append("")
    
    # Education
    if resume_data['education']:
        lines.append("EDUCATION")
        lines.append("-" * 70)
        for edu in resume_data['education']:
            lines.append(f"{edu['degree']} - {edu['institution']} ({edu['year']})")
            if edu.get('details'):
                lines.append(edu['details'])
            lines.append("")
        lines.append("")
    
    # Skills
    if resume_data['skills']:
        lines.append("SKILLS")
        lines.append("-" * 70)
        lines.append(" • ".join(resume_data['skills']))
        lines.append("")
    
    text_content = "\n".join(lines)
    
    response = HttpResponse(text_content, content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename="{user.username}_resume.txt"'
    return response




    
#export cover letter new
@login_required
def export_cover_letter(request):
    """
    Export cover letter as PDF, DOCX, or TXT.
    """
    format_type = request.GET.get("format", "pdf")
    
    # Get the cover letter
    cover_letter = Resume.objects.filter(
        user=request.user,
        status='draft',
        title__exact='Cover Letter Draft'
    ).first()

    if not cover_letter:
        return HttpResponse("No cover letter found. Please create one first.", status=404)

    # Extract cover letter data from conversation
    cover_data = extract_cover_letter_from_conversation(cover_letter)
    
    if not cover_data:
        return HttpResponse("No cover letter content found. Please complete the conversation.", status=404)

    # Generate the appropriate format
    if format_type == "pdf":
        return generate_cover_letter_pdf(request.user, cover_data)
    elif format_type == "docx":
        return generate_cover_letter_docx(request.user, cover_data)
    elif format_type == "txt":
        return generate_cover_letter_txt(request.user, cover_data)
    else:
        return HttpResponse("Invalid format", status=400)


def extract_cover_letter_from_conversation(cover_letter):
    """
    Extract structured cover letter data from conversation.
    """
    conversation = cover_letter.sections.get('conversation', [])
    
    if not conversation:
        return None
    
    # Get user responses
    user_responses = [msg['content'] for msg in conversation if msg.get('role') == 'user']
    ai_responses = [msg['content'] for msg in conversation if msg.get('role') == 'assistant']
    
    # Initialize cover letter data
    cover_data = {
        'name': cover_letter.user.get_full_name() or cover_letter.user.username,
        'email': cover_letter.user.email,
        'phone': '',
        'date': date.today().strftime("%B %d, %Y"),
        'company_name': 'Hiring Manager',
        'position': 'Position',
        'letter_body': ''
    }
    
    # Extract company and position from first user response
    if len(user_responses) > 0:
        first_response = user_responses[0].lower()
        
        # Try to extract company name
        if 'at' in first_response or 'with' in first_response:
            words = user_responses[0].split()
            for i, word in enumerate(words):
                if word.lower() in ['at', 'with'] and i + 1 < len(words):
                    cover_data['company_name'] = ' '.join(words[i+1:i+3])
                    break
        
        # Try to extract position
        for i, word in enumerate(user_responses[0].split()):
            if word.lower() in ['for', 'as']:
                cover_data['position'] = ' '.join(user_responses[0].split()[i+1:i+4])
                break
    
    # Get the final cover letter text
    # Look for the longest AI response (usually the final letter)
    final_letter = ""
    for response in reversed(ai_responses):
        if len(response) > 200:  # Substantial content
            lower = response.lower()
            # Check if it looks like a cover letter
            if any(phrase in lower for phrase in [
                'dear', 'sincerely', 'regards', 'position', 'experience', 'skills'
            ]):
                final_letter = response
                break
    
    # If no final letter found, combine last few AI responses
    if not final_letter and len(ai_responses) >= 2:
        final_letter = "\n\n".join(ai_responses[-2:])
    
    # If still nothing, create from user responses
    if not final_letter:
        final_letter = f"""Dear Hiring Manager,

I am writing to express my strong interest in the {cover_data['position']} position at {cover_data['company_name']}.

{' '.join(user_responses[:3])}

I am excited about the opportunity to contribute to your team and would welcome the chance to discuss how my skills align with your needs.

Thank you for considering my application.

Sincerely,
{cover_data['name']}"""
    
    cover_data['letter_body'] = final_letter
    
    return cover_data


def generate_cover_letter_pdf(user, cover_data):
    """
    Generate PDF cover letter.
    """

    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles for cover letter
    styles.add(ParagraphStyle(
        name='CoverLetterNormal',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        spaceAfter=12,
        alignment=TA_LEFT
    ))
    
    styles.add(ParagraphStyle(
        name='CoverLetterRight',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_RIGHT
    ))
    
    story = []
    
    # Your contact info (top right)
    contact_text = f"{cover_data['name']}<br/>{cover_data['email']}"
    if cover_data['phone']:
        contact_text += f"<br/>{cover_data['phone']}"
    story.append(Paragraph(contact_text, styles['CoverLetterRight']))
    story.append(Spacer(1, 0.3*inch))
    
    # Date
    story.append(Paragraph(cover_data['date'], styles['CoverLetterNormal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Recipient
    story.append(Paragraph(f"{cover_data['company_name']}<br/>{cover_data['position']}", styles['CoverLetterNormal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Letter body
    # Split into paragraphs
    paragraphs = cover_data['letter_body'].split('\n\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para.strip().replace('\n', '<br/>'), styles['CoverLetterNormal']))
            story.append(Spacer(1, 0.15*inch))
    
    doc.build(story)
    
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{user.username}_cover_letter.pdf"'
    return response


def generate_cover_letter_docx(user, cover_data):
    """
    Generate DOCX cover letter.
    """
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Your contact info (right aligned)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact_run = contact_para.add_run(f"{cover_data['name']}\n{cover_data['email']}")
    contact_run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacer
    
    # Date
    doc.add_paragraph(cover_data['date'])
    doc.add_paragraph()  # Spacer
    
    # Recipient
    doc.add_paragraph(cover_data['company_name'])
    doc.add_paragraph(cover_data['position'])
    doc.add_paragraph()  # Spacer
    
    # Letter body
    paragraphs = cover_data['letter_body'].split('\n\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.space_after = Pt(12)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{user.username}_cover_letter.docx"'
    return response


def generate_cover_letter_txt(user, cover_data):
    """
    Generate TXT cover letter.
    """
    lines = []
    
    # Your contact info
    lines.append(cover_data['name'])
    lines.append(cover_data['email'])
    if cover_data['phone']:
        lines.append(cover_data['phone'])
    lines.append("")
    lines.append("")
    
    # Date
    lines.append(cover_data['date'])
    lines.append("")
    lines.append("")
    
    # Recipient
    lines.append(cover_data['company_name'])
    lines.append(cover_data['position'])
    lines.append("")
    lines.append("")
    
    # Letter body
    lines.append(cover_data['letter_body'])
    
    text_content = "\n".join(lines)
    
    response = HttpResponse(text_content, content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename="{user.username}_cover_letter.txt"'
    return response